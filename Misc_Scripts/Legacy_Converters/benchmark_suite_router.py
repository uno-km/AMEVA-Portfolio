import os
import csv
import time
import re
import json
import sqlite3
import tempfile
import threading
import psutil
import subprocess
import requests
import pandas as pd
from typing import List, Dict, Any
from datetime import datetime
from fastapi import APIRouter, BackgroundTasks, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from core.constants import OLLAMA_BASE_URL, LLAMA_CPP_HOST, LLAMA_CPP_PORT
from core.judge_service import JudgeService
from core.prompt_utils import PromptFactory, get_stop_tokens
from backend.state import state, CONFIG_PATH
from backend.routers.logs import broadcaster
from backend.database import get_db_connection

router = APIRouter()

# ─────────────────────────────────────────────────────────────────────────────
# Schemas
# ─────────────────────────────────────────────────────────────────────────────

class BootConfigSchema(BaseModel):
    engine: str = "OLM"
    cpu_cores: float = 2.0
    ram_mb: int = 4096
    gpu_layers: int = 0
    model_name: str = ""

class StressOptionsSchema(BaseModel):
    threads: int = 4
    n_ctx: int = 2048
    iterations: int = 1
    temperature: float = 0.1
    top_k: int = 40
    top_p: float = 0.95
    repeat_penalty: float = 1.1
    system_prompt: str = "You are a professional benchmark assistant. Answer precisely and concisely."
    judge_model: str = "exaone3.5:7.8b"

class RunBenchmarkRequest(BaseModel):
    boot_config: BootConfigSchema
    stress_config: StressOptionsSchema
    run_mode: str = "Inference Mode (Default)"

class SyncBenchmarkRequest(BaseModel):
    boot_config: BootConfigSchema
    stress_config: StressOptionsSchema
    run_mode: str = "Inference Mode (Default)"
    harness_tasks: List["TaskSchema"] = None

class ChatRequest(BaseModel):
    prompt: str
    boot_config: BootConfigSchema
    stress_config: StressOptionsSchema

class TaskSchema(BaseModel):
    task_id: str
    prompt: str
    expected_regex: str = ""
    eval_type: str = "llm_judge"

class ExportWordRequest(BaseModel):
    run_id: int
    use_llm_summary: bool = True

# ─────────────────────────────────────────────────────────────────────────────
# Config API
# ─────────────────────────────────────────────────────────────────────────────

@router.get("/api/config")
def get_config_endpoint():
    return state.config_data

@router.post("/api/config")
def update_config_endpoint(updates: dict):
    state.save_config(updates)
    # config.json 저장에 맞춰 state 변수 동기화
    if "default_judge_model" in updates:
        state.stress_config.judge_model = updates["default_judge_model"]
    return {"status": "saved", "config": state.config_data}

# ─────────────────────────────────────────────────────────────────────────────
# Power Tracker (Threaded)
# ─────────────────────────────────────────────────────────────────────────────

class PowerTracker:
    def __init__(self, has_nvidia: bool = False):
        self.is_running = True
        self.has_nvidia = has_nvidia
        self.power_history = []
        self.thread = None

    def start(self):
        self.is_running = True
        self.thread = threading.Thread(target=self._run)
        self.thread.daemon = True
        self.thread.start()

    def _run(self):
        while self.is_running:
            try:
                watts = 0.0
                if self.has_nvidia:
                    try:
                        proc = subprocess.Popen(
                            ["nvidia-smi", "--query-gpu=power.draw", "--format=csv,noheader,nounits"],
                            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                            text=True, creationflags=subprocess.CREATE_NO_WINDOW
                        )
                        stdout, _ = proc.communicate(timeout=0.1)
                        lines = stdout.strip().split('\n')
                        if lines and lines[0]:
                            watts += float(lines[0])
                    except:
                        pass
                
                cpu_p = psutil.cpu_percent()
                watts += 5.0 + (cpu_p * 0.6) 
                self.power_history.append(watts)
            except Exception:
                pass
            time.sleep(0.2)

    def stop(self):
        self.is_running = False
        if self.thread:
            self.thread.join(timeout=1.0)

    def get_average_watts(self) -> float:
        if not self.power_history:
            return 0.0
        return sum(self.power_history) / len(self.power_history)

# ─────────────────────────────────────────────────────────────────────────────
# Helper Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _ts() -> str:
    now = datetime.now()
    return f"[{now.strftime('%H:%M:%S')}.{now.microsecond // 1000:03d}]"

def boot_worker(config_dict: dict):
    state.boot_status = "BOOTING"
    state.boot_message = "BOOTING..."
    
    state.engine.set_logger(lambda msg: broadcaster.log(f"{_ts()} {msg}", "sys"))
    
    success, msg = state.engine.boot_matrix(config_dict)
    
    if success:
        state.boot_status = "ONLINE"
        state.boot_message = msg
        state.last_booted_model = config_dict.get("model_name", "")
        broadcaster.log(f"✅ 부팅 완료: {msg}", "sys")
        
        # 마지막 성공 세션 캐시 설정 저장
        state.save_config({
            "last_used_engine": config_dict.get("engine", "OLM"),
            "last_used_cores": config_dict.get("cpu_cores", 2.0),
            "last_used_ram": config_dict.get("ram_mb", 4096),
            "last_used_gpu_layers": config_dict.get("gpu_layers", 0)
        })
    else:
        state.boot_status = "ERROR"
        state.boot_message = f"부팅 실패: {msg}"
        broadcaster.log(f"❌ 부팅 실패: {msg}", "sys")

# ─────────────────────────────────────────────────────────────────────────────
# Endpoints
# ─────────────────────────────────────────────────────────────────────────────

@router.get("/api/session/status")
def get_session_status():
    return {
        "boot_status": state.boot_status,
        "boot_message": state.boot_message,
        "last_booted_model": state.last_booted_model,
        "benchmark_running": state.active_benchmark_running,
        "chat_running": state.active_chat_running
    }

@router.post("/api/session/boot")
def boot_session(req: BootConfigSchema):
    if state.active_benchmark_running or state.active_chat_running:
        raise HTTPException(status_code=400, detail="Benchmark or chat inference is currently active.")
        
    config_dict = {
        "engine": req.engine,
        "cpu_cores": req.cpu_cores,
        "ram_mb": req.ram_mb,
        "gpu_layers": req.gpu_layers,
        "model_name": req.model_name
    }
    
    state.boot_config = req
    state.session.boot_config = req
    
    t = threading.Thread(target=boot_worker, args=(config_dict,))
    t.start()
    
    return {"status": "booting"}

@router.post("/api/session/shutdown")
def shutdown_session():
    if state.active_benchmark_running or state.active_chat_running:
        raise HTTPException(status_code=400, detail="Cannot shutdown while task is running.")
        
    broadcaster.log("📢 시스템 리부트 시퀀스: 자원 완전 반납 중...", "sys")
    state.engine.shutdown()
    
    state.last_booted_model = ""
    state.boot_status = "OFFLINE"
    state.boot_message = "READY"
    
    return {"status": "shutdown_done"}

# ── Harness CRUD (SQLite 연동) ──

@router.get("/api/harness")
def get_harness():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT task_id, prompt, expected_regex, eval_type FROM harness_tasks;")
    rows = cursor.fetchall()
    conn.close()
    
    # 만약 비어있다면 기초 마이그레이션 확인
    if not rows:
        from backend.database import db_init
        db_init()
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT task_id, prompt, expected_regex, eval_type FROM harness_tasks;")
        rows = cursor.fetchall()
        conn.close()
        
    return [dict(r) for r in rows]

@router.post("/api/harness")
def save_harness(tasks: List[TaskSchema]):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # 데이터 클리어 후 재삽입
        cursor.execute("DELETE FROM harness_tasks;")
        for t in tasks:
            cursor.execute("""
            INSERT INTO harness_tasks (task_id, prompt, expected_regex, eval_type)
            VALUES (?, ?, ?, ?);
            """, (t.task_id, t.prompt, t.expected_regex, t.eval_type))
        conn.commit()
        return {"status": "saved", "count": len(tasks)}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to write harness tasks: {e}")
    finally:
        conn.close()

# ── Reports ──

@router.get("/api/reports")
def get_reports(n: int = 50):
    conn = get_db_connection()
    cursor = conn.cursor()
    # 개별 태스크 평균 점수 및 TPS를 포함한 런 리스트 쿼리
    cursor.execute("""
    SELECT r.id, r.timestamp, r.model_name, r.engine_type, r.run_mode, 
           r.cpu_cores, r.ram_mb, r.gpu_layers, r.threads, r.n_ctx, r.judge_model,
           COUNT(s.id) as task_count,
           AVG(CASE WHEN s.judge_score NOT IN ('PASS (Regex)', 'FAIL (Regex)', 'N/A') THEN CAST(s.judge_score AS REAL) ELSE NULL END) as avg_score,
           AVG(s.tps) as avg_tps,
           AVG(s.ttft_ms) as avg_ttft,
           AVG(s.avg_gpu_w) as avg_power
    FROM benchmark_runs r
    LEFT JOIN benchmark_results s ON r.id = s.run_id
    GROUP BY r.id
    ORDER BY r.id DESC
    LIMIT ?;
    """, (n,))
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

@router.get("/api/reports/{run_id}")
def get_report_detail(run_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. 런 메타데이터
    cursor.execute("SELECT * FROM benchmark_runs WHERE id = ?;", (run_id,))
    run_row = cursor.fetchone()
    if not run_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Benchmark run not found")
        
    # 2. 개별 상세 결과
    cursor.execute("SELECT * FROM benchmark_results WHERE run_id = ?;", (run_id,))
    result_rows = cursor.fetchall()
    conn.close()
    
    return {
        "run": dict(run_row),
        "results": [dict(r) for r in result_rows]
    }

@router.delete("/api/reports/{run_id}")
def delete_report(run_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM benchmark_runs WHERE id = ?;", (run_id,))
        conn.commit()
        return {"status": "deleted"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ─────────────────────────────────────────────────────────────────────────────
# Execution Workers
# ─────────────────────────────────────────────────────────────────────────────

def run_benchmark_worker(req: RunBenchmarkRequest):
    state.active_benchmark_running = True
    broadcaster.log(f"🚀 벤치마크 가동 시퀀스 시작 (모드: {req.run_mode})", "bench")
    
    try:
        # 1. Smart SWAP 체크
        if state.last_booted_model != req.boot_config.model_name:
            broadcaster.log(f"🔄 스왑 필요 감지: {state.last_booted_model} -> {req.boot_config.model_name}", "sys")
            config_dict = {
                "engine": req.boot_config.engine,
                "cpu_cores": req.boot_config.cpu_cores,
                "ram_mb": req.boot_config.ram_mb,
                "gpu_layers": req.boot_config.gpu_layers,
                "model_name": req.boot_config.model_name
            }
            state.engine.set_logger(lambda msg: broadcaster.log(f"{_ts()} {msg}", "sys"))
            success, msg = state.engine.boot_matrix(config_dict)
            if not success:
                state.boot_status = "ERROR"
                state.boot_message = f"스왑 실패: {msg}"
                broadcaster.log(f"❌ 스왑 실패: {msg}", "sys")
                return
            state.boot_status = "ONLINE"
            state.boot_message = msg
            state.last_booted_model = req.boot_config.model_name

        # 2. 실행 분기
        if "Stress" in req.run_mode or "Hard" in req.run_mode:
            _run_stress_mode(req)
        else:
            _run_inference_mode(req)
            
    except Exception as e:
        broadcaster.log(f"[FATAL] 벤치마크 수행 중 예외 발생: {e}", "sys")
    finally:
        state.active_benchmark_running = False

def _run_stress_mode(req: RunBenchmarkRequest) -> tuple[int, list[dict]]:
    broadcaster.log("LLAMA-BENCH 스트레스 테스트 시작", "bench")
    results = []

    pw_tracker = PowerTracker()
    if "Efficiency" in req.run_mode:
        pw_tracker.start()
    start_time = time.time()

    opts = {
        'threads': req.stress_config.threads,
        'n_ctx':   req.stress_config.n_ctx
    }
    
    bench_data = state.engine.run_llama_bench(req.boot_config.model_name, opts)

    if "Efficiency" in req.run_mode:
        pw_tracker.stop()

    avg_watts = pw_tracker.get_average_watts()
    
    if not bench_data:
        broadcaster.log("[에러] 스트레스 테스트에서 데이터를 반환하지 못했습니다.", "bench")
        return 0, []

    # 데이터베이스 삽입
    conn = get_db_connection()
    cursor = conn.cursor()
    run_id = 0
    try:
        cursor.execute("""
        INSERT INTO benchmark_runs (
            timestamp, model_name, engine_type, run_mode, cpu_cores, ram_mb, gpu_layers, 
            threads, n_ctx, temperature, repeat_penalty, system_prompt, judge_model
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);
        """, (
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"), req.boot_config.model_name, req.boot_config.engine,
            req.run_mode, req.boot_config.cpu_cores, req.boot_config.ram_mb, req.boot_config.gpu_layers,
            req.stress_config.threads, req.stress_config.n_ctx, req.stress_config.temperature,
            req.stress_config.repeat_penalty, req.stress_config.system_prompt, req.stress_config.judge_model
        ))
        run_id = cursor.lastrowid

        for item in bench_data:
            cursor.execute("""
            INSERT INTO benchmark_results (
                run_id, task_name, category, prompt_text, response_text, ttft_ms, 
                prompt_eval_ms_t, avg_gpu_w, tokens_per_joule, e2e_latency_sec, tps, 
                peak_vram_mb, system_load, warm_cold_tag, sampling_time_ms, judge_score, judge_reason
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);
            """, (
                run_id, "Llama-Bench", "STRESS", "N/A", "N/A", 0.0, 0.0, avg_watts,
                round(item.get('t/s', 0)/avg_watts, 3) if avg_watts > 0 else 0,
                0.0, round(item.get('t/s', 0), 2), 0.0, "STRESS", "STRESS", 0.0, "N/A", "N/A"
            ))
            results.append({
                "task_name": "Llama-Bench",
                "category": "STRESS",
                "prompt_text": "N/A",
                "response_text": "N/A",
                "ttft_ms": 0.0,
                "tps": round(item.get('t/s', 0), 2),
                "avg_gpu_w": round(avg_watts, 2),
                "tokens_per_joule": round(item.get('t/s', 0)/avg_watts, 3) if avg_watts > 0 else 0,
                "judge_score": "N/A",
                "judge_reason": "N/A"
            })
        conn.commit()
        broadcaster.log(f"📊 {len(bench_data)}건 스트레스 결과가 SQLite DB에 저장되었습니다.", "bench")
    except Exception as e:
        conn.rollback()
        broadcaster.log(f"❌ DB 저장 에러: {e}", "bench")
    finally:
        conn.close()
    
    # 자원 반납
    state.engine.shutdown()
    state.last_booted_model = ""
    state.boot_status = "OFFLINE"
    state.boot_message = "READY"
    broadcaster.log("✓ 벤치마크 완료 및 엔진 종료.", "bench")
    return run_id, results
def _run_inference_mode(req: RunBenchmarkRequest, tasks_list: list[dict] = None) -> tuple[int, list[dict]]:
    # 하네스 데이터 로드
    if tasks_list is not None:
        dataset = tasks_list
    else:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT task_id, prompt, expected_regex, eval_type FROM harness_tasks;")
        dataset = [dict(r) for r in cursor.fetchall()]
        conn.close()

    if not dataset:
        broadcaster.log("[에러] 하네스 태스크가 존재하지 않습니다. 먼저 하네스 매니저에서 추가해 주세요.", "bench")
        return 0, []

    broadcaster.log(f"추론 모드 시작 – 하네스 태스크 {len(dataset)}개", "bench")
    results = []

    engine_type = req.boot_config.engine
    model_name  = req.boot_config.model_name

    if engine_type == "OLM":
        try:
            OllamaClient.pull_model_stream(model_name)
        except:
            pass
            
    from models.hardware import HardwareService
    specs = HardwareService.detect_capabilities()
    has_nv = specs.has_nvidia

    for idx, task in enumerate(dataset):
        task_id = task.get('task_id', '?')
        eval_type = task.get('eval_type', 'llm_judge')
        raw_prompt = task.get('prompt', 'Hello')

        broadcaster.log(f"====== ### [Task {idx+1}/{len(dataset)}] {task_id} ======", "bench")
        broadcaster.log(f"====== ### EVAL TYPE: {eval_type}", "bench")
        broadcaster.log(f"====== ### QUESTION: {raw_prompt}", "bench")
        broadcaster.log(f"[STATUS] Running inference on {model_name}...", "bench")

        cat_name = "General"
        if "-" in task_id:
            cat_name = task_id.split("-")[0]
            
        broadcaster.log(f"\n[INFO] AI가 '{task_id}' 문제를 분석 중입니다... (TTFT 측정 중)\n", "chunk")

        pw_tracker = PowerTracker(has_nvidia=has_nv)
        pw_tracker.start()

        start_time = time.time()
        text_acc = ""
        ttft = 0
        prompt_ms_per_t = 0
        sample_ms = 0
        tok_count = 0

        formatted_prompt = PromptFactory.wrap(raw_prompt, model_name, req.stress_config.system_prompt)
        stop_tokens = get_stop_tokens(model_name)

        if engine_type == "OLM":
            payload = {
                "model": model_name,
                "prompt": formatted_prompt,
                "stream": True,
                "options": {
                    "num_predict": 200,
                    "num_thread":  req.stress_config.threads,
                    "temperature": req.stress_config.temperature,
                    "top_k": req.stress_config.top_k,
                    "top_p": req.stress_config.top_p,
                    "repeat_penalty": req.stress_config.repeat_penalty,
                    "stop": stop_tokens
                }
            }
            url = f"{OLLAMA_BASE_URL}/api/generate"
        else:
            payload = {
                "prompt": formatted_prompt,
                "stream": True,
                "n_predict": 200,
                "temperature": req.stress_config.temperature,
                "top_k": req.stress_config.top_k,
                "top_p": req.stress_config.top_p,
                "repeat_penalty": req.stress_config.repeat_penalty,
                "stop": stop_tokens
            }
            url = f"http://{LLAMA_CPP_HOST}:{LLAMA_CPP_PORT}/completion"

        try:
            if engine_type == "BIT":
                if state.engine.container:
                    from core.constants import get_filename_by_id
                    model_file = get_filename_by_id(model_name)
                    cmd = f'python3 run_inference.py -m /vault/{model_file} -p "{formatted_prompt}" -n 200'
                    exit_code, output = state.engine.container.exec_run(cmd)
                    text_acc = output.decode('utf-8', errors='replace')
                    if "Answer:" in text_acc:
                        text_acc = text_acc.split("Answer:")[-1].strip()
                    tok_count = len(text_acc.split())
                    if ttft == 0: ttft = (time.time() - start_time) * 1000
                    broadcaster.log(text_acc, "chunk")
            elif engine_type == "OLM":
                resp = requests.post(url, json=payload, stream=True, timeout=60)
                resp.raise_for_status()
                for raw_line in resp.iter_lines():
                    if not raw_line: continue
                    try:
                        data = json.loads(raw_line.decode('utf-8'))
                    except (json.JSONDecodeError, UnicodeDecodeError): continue
                    if ttft == 0: ttft = (time.time() - start_time) * 1000
                    token = data.get('response', '')
                    if token:
                        text_acc += token
                        tok_count += 1
                        broadcaster.log(token, "chunk")
                    if data.get('done'):
                        eval_count = data.get('eval_count', 0)
                        eval_dur_ns = data.get('eval_duration', 0)
                        if eval_count > 0 and eval_dur_ns > 0:
                            sample_ms = eval_dur_ns / 1e6
                        break
            else:
                resp = requests.post(url, json=payload, stream=True, timeout=60)
                resp.raise_for_status()
                buffer = b""
                for chunk in resp.iter_content(chunk_size=512):
                    if not chunk: continue
                    buffer += chunk
                    while b"\n" in buffer:
                        line_bytes, buffer = buffer.split(b"\n", 1)
                        line = line_bytes.decode('utf-8', errors='replace').strip()
                        if not line.startswith("data:"): continue
                        payload_str = line[5:].strip()
                        if payload_str == "[DONE]": break
                        try:
                            data = json.loads(payload_str)
                        except json.JSONDecodeError: continue
                        if ttft == 0: ttft = (time.time() - start_time) * 1000
                        token = data.get('content', '')
                        if token:
                            text_acc += token
                            tok_count += 1
                            broadcaster.log(token, "chunk")
                        if data.get('stop'):
                            t_info = data.get('timings', {})
                            prompt_n = t_info.get('prompt_n', 1)
                            p_ms = t_info.get('prompt_ms', 0)
                            prompt_ms_per_t = round(p_ms / prompt_n, 2) if prompt_n > 0 else 0
                            sample_ms = t_info.get('predicted_ms', 0)
                            break
        except Exception as e:
            broadcaster.log(f"[에러] 추론 엔진 통신 실패: {e}", "bench")

        duration = time.time() - start_time
        if ttft == 0: ttft = duration * 1000
        
        broadcaster.log(f"\n[DONE] 생성 완료. (TTFT: {ttft:.1f}ms / TPS: {tok_count/duration if duration>0 else 0:.2f})\n", "chunk")
        
        pw_tracker.stop()

        eval_type = task.get('eval_type', 'llm_judge')
        score = "N/A"
        if eval_type == 'regex':
            pattern = task.get('expected_regex', '')
            if pattern and re.search(pattern, text_acc):
                score = "PASS (Regex)"
            else:
                score = "FAIL (Regex)"

        avg_watts = pw_tracker.get_average_watts()
        tps_val = round(tok_count / duration, 2) if duration > 0 else 0

        broadcaster.log(f"Task 완료 | TPS: {tps_val} | TTFT: {ttft:.1f}ms | {avg_watts:.1f}W", "bench")
        broadcaster.log(f"✓ [{cat_name}] {task.get('task_id','?')}  |  Judge: {score}  |  {duration:.2f}s  |  {tps_val} t/s", "bench")

        results.append({
            "task_name":         task.get("task_id"),
            "category":          cat_name,
            "prompt_text":        task.get('prompt', ''),
            "response_text":    text_acc,
            "ttft_ms":          round(ttft, 1),
            "prompt_eval_ms_t": prompt_ms_per_t,
            "avg_gpu_w":          round(avg_watts, 2),
            "tokens_per_joule":   round(tps_val / avg_watts, 3) if avg_watts > 0 else 0,
            "e2e_latency_sec":        round(duration, 2),
            "tps":   tps_val,
            "peak_vram_mb":       0,
            "system_load":        "INFERENCE",
            "warm_cold_tag":      "WARM",
            "sampling_time_ms": round(sample_ms, 2),
            "judge_score":        score,
            "judge_reason":       "N/A",
            "eval_type":          eval_type
        })

    broadcaster.log("✓ 벤치마크 추론 시퀀스 완료.", "bench")
    
    # 판정 전 메인 엔진 리소스 명시적 해제
    broadcaster.log("⚙️  판정 전 리소스 최적화: 메인 엔진 언로드 시퀀스...", "sys")
    state.engine.shutdown()
    state.last_booted_model = ""
    state.boot_status = "OFFLINE"
    state.boot_message = "READY"
    time.sleep(1.0)

    final_scores = []
    total_judge = len(results)
    broadcaster.log(f"🧠 판정관 가동: {state.stress_config.judge_model} (총 {total_judge}개 채점 대상)", "bench")
    
    judge_idx = 0
    try:
        for res in results:
            judge_idx += 1
            broadcaster.log(f"🧠 [판정 진행률: {judge_idx}/{total_judge}] '{res['task_name']}' 채점 중...", "bench")
            score_data = JudgeService.call_llm_judge(
                res["prompt_text"], 
                res["response_text"], 
                state.stress_config,
                chunk_callback=lambda tok: broadcaster.log(tok, "chunk")
            )
            
            llm_score = str(score_data.get("score", 0))
            llm_reason = score_data.get("reason", "No reason provided.")
            
            if res.get("eval_type") == "regex":
                res["judge_reason"] = f"[Regex Result: {res['judge_score']}] " + llm_reason
                res["judge_score"] = llm_score
            else:
                res["judge_score"]  = llm_score
                res["judge_reason"] = llm_reason
                
            broadcaster.log(f"   └ [판정관 의견]: {res['judge_reason']}", "bench")
            
            try:
                final_scores.append(int(llm_score))
            except:
                final_scores.append(0)
                
        broadcaster.log("✅ 모든 판정이 성공적으로 완료되었습니다.", "bench")
    except Exception as e:
        broadcaster.log(f"❌ 판정 수행 중 오류: {e}", "bench")

    avg_score = sum(final_scores)/len(final_scores) if final_scores else 0
    
    broadcaster.log("\n" + "="*50, "bench")
    broadcaster.log("🏆 [AMEVA] 최종 벤치마크 리포트 (EXAONE 3.5 기준)", "bench")
    broadcaster.log("="*50, "bench")
    broadcaster.log(f"{'CATEGORY':<15} | {'SCORE':<10} | {'STATUS'}", "bench")
    broadcaster.log("-" * 50, "bench")
    
    cat_scores = {}
    for r in results:
        cat = r.get("category", "General")
        try:
            val = float(r["judge_score"])
            if cat not in cat_scores: cat_scores[cat] = []
            cat_scores[cat].append(val)
        except:
            pass
        
    for cat, scores in cat_scores.items():
        if scores:
            c_avg = sum(scores) / len(scores)
            status_text = f"{c_avg:.2f}"
        else:
            status_text = "N/A"
        broadcaster.log(f"{cat:<15} | {status_text:<10} | OK", "bench")
        
    broadcaster.log("-" * 50, "bench")
    broadcaster.log(f"⭐ TOTAL AVERAGE: {avg_score:.2f} / 10.0", "bench")
    broadcaster.log("="*50, "bench")

    # DB 저장
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
        INSERT INTO benchmark_runs (
            timestamp, model_name, engine_type, run_mode, cpu_cores, ram_mb, gpu_layers, 
            threads, n_ctx, temperature, repeat_penalty, system_prompt, judge_model
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);
        """, (
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"), model_name, engine_type,
            req.run_mode, req.boot_config.cpu_cores, req.boot_config.ram_mb, req.boot_config.gpu_layers,
            req.stress_config.threads, req.stress_config.n_ctx, req.stress_config.temperature,
            req.stress_config.repeat_penalty, req.stress_config.system_prompt, state.stress_config.judge_model
        ))
        run_id = cursor.lastrowid
        
        for res in results:
            cursor.execute("""
            INSERT INTO benchmark_results (
                run_id, task_name, category, prompt_text, response_text, ttft_ms, 
                prompt_eval_ms_t, avg_gpu_w, tokens_per_joule, e2e_latency_sec, tps, 
                peak_vram_mb, system_load, warm_cold_tag, sampling_time_ms, judge_score, judge_reason
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);
            """, (
                run_id, res["task_name"], res["category"], res["prompt_text"], res["response_text"],
                res["ttft_ms"], res["prompt_eval_ms_t"], res["avg_gpu_w"], res["tokens_per_joule"],
                res["e2e_latency_sec"], res["tps"], res["peak_vram_mb"], "INFERENCE", "WARM",
                res["sampling_time_ms"], res["judge_score"], res["judge_reason"]
            ))
        conn.commit()
        broadcaster.log("📊 결과 저장 완료 및 시퀀스 리셋.", "bench")
    except Exception as e:
        conn.rollback()
        broadcaster.log(f"❌ DB 저장 에러: {e}", "bench")
    finally:
        conn.close()

@router.post("/api/benchmark/run")
def run_benchmark(req: RunBenchmarkRequest, background_tasks: BackgroundTasks):
    if state.active_benchmark_running or state.active_chat_running:
        raise HTTPException(status_code=400, detail="A benchmark or chat session is already running.")
    
    background_tasks.add_task(run_benchmark_worker, req)
    return {"status": "started"}

# ── Chat Worker & Endpoint ──

def execute_chat_logic(req: ChatRequest) -> tuple[int, dict]:
    engine_type = req.boot_config.engine
    model_name = req.boot_config.model_name
    
    formatted_prompt = PromptFactory.wrap(req.prompt, model_name, req.stress_config.system_prompt)
    stop_tokens = get_stop_tokens(model_name)

    broadcaster.log(f"[CHAT_MOD] 채팅 벤치마크 시작 – 모델: {model_name}", "sys")
    
    if state.last_booted_model != model_name:
        broadcaster.log(f"🔄 스왑 필요 감지: {state.last_booted_model} -> {model_name}", "sys")
        config_dict = {
            "engine": engine_type,
            "cpu_cores": req.boot_config.cpu_cores,
            "ram_mb": req.boot_config.ram_mb,
            "gpu_layers": req.boot_config.gpu_layers,
            "model_name": model_name
        }
        state.engine.set_logger(lambda msg: broadcaster.log(f"{_ts()} {msg}", "sys"))
        success, msg = state.engine.boot_matrix(config_dict)
        if not success:
            state.boot_status = "ERROR"
            state.boot_message = f"스왑 실패: {msg}"
            broadcaster.log(f"❌ 스왑 실패: {msg}", "sys")
            raise RuntimeError(f"Engine swap failed: {msg}")
        state.boot_status = "ONLINE"
        state.boot_message = msg
        state.last_booted_model = model_name

    sc = req.stress_config
    if engine_type == "OLM":
        url = f"{OLLAMA_BASE_URL}/api/generate"
        payload = {
            "model":   model_name,
            "prompt":  formatted_prompt,
            "stream":  True,
            "options": {
                "num_thread": sc.threads,
                "temperature": sc.temperature,
                "top_k": sc.top_k,
                "top_p": sc.top_p,
                "repeat_penalty": sc.repeat_penalty,
                "stop": stop_tokens
            },
        }
    else:
        url = f"http://{LLAMA_CPP_HOST}:{LLAMA_CPP_PORT}/completion"
        payload = {
            "prompt":    formatted_prompt,
            "stream":    True,
            "n_predict": 512,
            "temperature": sc.temperature,
            "top_k": sc.top_k,
            "top_p": sc.top_p,
            "repeat_penalty": sc.repeat_penalty,
            "stop": stop_tokens
        }

    text_acc         = ""
    ttft             = 0
    prompt_ms_per_t  = 0
    sample_ms        = 0
    tok_count        = 0
    start_time       = time.time()

    try:
        if engine_type == "BIT":
            if state.engine.container:
                from core.constants import get_filename_by_id
                model_file = get_filename_by_id(model_name)
                cmd = f'python3 run_inference.py -m /vault/{model_file} -p "{formatted_prompt}" -n 200'
                exit_code, output = state.engine.container.exec_run(cmd)
                text_acc = output.decode('utf-8', errors='replace')
                if "Answer:" in text_acc:
                    text_acc = text_acc.split("Answer:")[-1].strip()
                tok_count = len(text_acc.split())
                if ttft == 0: ttft = (time.time() - start_time) * 1000
                broadcaster.log(text_acc, "chunk")
        elif engine_type == "OLM":
            resp = requests.post(url, json=payload, stream=True, timeout=60)
            resp.raise_for_status()
            for raw_line in resp.iter_lines():
                if not raw_line: continue
                try:
                    data = json.loads(raw_line.decode('utf-8'))
                except (json.JSONDecodeError, UnicodeDecodeError): continue
                if ttft == 0: ttft = (time.time() - start_time) * 1000
                token = data.get('response', '')
                if token:
                    text_acc += token
                    tok_count += 1
                    broadcaster.log(token, "chunk")
                if data.get('done'):
                    eval_count = data.get('eval_count', 0)
                    eval_dur_ns = data.get('eval_duration', 0)
                    if eval_count > 0 and eval_dur_ns > 0:
                        sample_ms = eval_dur_ns / 1e6
                    break
        else:
            resp = requests.post(url, json=payload, stream=True, timeout=60)
            resp.raise_for_status()
            buffer = b""
            for chunk in resp.iter_content(chunk_size=512):
                if not chunk: continue
                buffer += chunk
                while b"\n" in buffer:
                    line_bytes, buffer = buffer.split(b"\n", 1)
                    line = line_bytes.decode('utf-8', errors='replace').strip()
                    if not line.startswith("data:"): continue
                    payload_str = line[5:].strip()
                    if payload_str == "[DONE]": break
                    try:
                        data = json.loads(payload_str)
                    except json.JSONDecodeError: continue
                    if ttft == 0: ttft = (time.time() - start_time) * 1000
                    token = data.get('content', '')
                    if token:
                        text_acc += token
                        tok_count += 1
                        broadcaster.log(token, "chunk")
                    if data.get('stop'):
                        t = data.get('timings', {})
                        pn = t.get('prompt_n', 0)
                        pm = t.get('prompt_ms', 0)
                        prompt_ms_per_t = round(pm / pn, 2) if pn > 0 else 0
                        sample_ms = t.get('predicted_ms', 0)
                        break
    except Exception as e:
        broadcaster.log(f"[CHAT_MOD] 오류: {e}", "sys")
        raise e

    duration = time.time() - start_time
    if ttft == 0: ttft = duration * 1000
    tps_val = round(tok_count / duration, 2) if duration > 0 else 0

    broadcaster.log(f"[CHAT_MOD] 완료 | TPS: {tps_val} | TTFT: {ttft:.1f}ms | {duration:.2f}s", "sys")

    result = {
        "Model_Hash":          model_name,
        "Quant_Method":        "N/A",
        "Context_Size":        req.stress_config.n_ctx,
        "Thread_Config":       req.stress_config.threads,
        "Prompt_Text":         req.prompt,
        "Prompt_Response":     text_acc,
        "TTFT (ms)":           round(ttft, 1),
        "Prompt_Eval (ms/t)":  prompt_ms_per_t,
        "Avg_GPU_W":           0,
        "Tokens_per_Joule":    0,
        "E2E_Latency":         round(duration, 2),
        "Generation (t/s)":    tps_val,
        "Peak_VRAM_MB":        0,
        "System_Load":         "[CHAT_MOD]",
        "Warm/Cold_Tag":       "CHAT",
        "Sampling_Time (ms)":  round(sample_ms, 2),
        "Judge_Score":         "N/A",
        "Judge_Reason":        "N/A"
    }

    if req.stress_config.judge_model:
        broadcaster.log(f"🧠 판정관 호출 중: {req.stress_config.judge_model}", "sys")
        try:
            score_data = JudgeService.call_llm_judge(
                req.prompt, 
                text_acc, 
                req.stress_config,
                chunk_callback=lambda tok: broadcaster.log(tok, "chunk")
            )
            result["Judge_Score"] = str(score_data.get("score", 0))
            result["Judge_Reason"] = score_data.get("reason", "")
            broadcaster.log(f"🏆 채팅 판정 완료: {result['Judge_Score']}/10", "sys")
            broadcaster.log("✅ 모든 판정이 성공적으로 완료되었습니다.", "sys")
        except Exception as e:
            broadcaster.log(f"❌ 판정관 호출 중 치명적 오류: {e}", "sys")
            result["Judge_Score"] = "0"
            result["Judge_Reason"] = f"Error: {e}"

    # DB 저장
    conn = get_db_connection()
    cursor = conn.cursor()
    run_id = 0
    try:
        cursor.execute("""
        INSERT INTO benchmark_runs (
            timestamp, model_name, engine_type, run_mode, cpu_cores, ram_mb, gpu_layers, 
            threads, n_ctx, temperature, repeat_penalty, system_prompt, judge_model
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);
        """, (
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"), model_name, engine_type,
            "[CHAT_MOD]", req.boot_config.cpu_cores, req.boot_config.ram_mb, req.boot_config.gpu_layers,
            req.stress_config.threads, req.stress_config.n_ctx, req.stress_config.temperature,
            req.stress_config.repeat_penalty, req.stress_config.system_prompt, req.stress_config.judge_model
        ))
        run_id = cursor.lastrowid
        
        cursor.execute("""
        INSERT INTO benchmark_results (
            run_id, task_name, category, prompt_text, response_text, ttft_ms, 
            prompt_eval_ms_t, avg_gpu_w, tokens_per_joule, e2e_latency_sec, tps, 
            peak_vram_mb, system_load, warm_cold_tag, sampling_time_ms, judge_score, judge_reason
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);
        """, (
            run_id, "Single Chat", "CHAT", result["Prompt_Text"], result["Prompt_Response"],
            result["TTFT (ms)"], result["Prompt_Eval (ms/t)"], result["Avg_GPU_W"], result["Tokens_per_Joule"],
            result["E2E_Latency"], result["Generation (t/s)"], result["Peak_VRAM_MB"], "CHAT", "CHAT",
            result["Sampling_Time (ms)"], result["Judge_Score"], result["Judge_Reason"]
        ))
        conn.commit()
        broadcaster.log("[CHAT_MOD] SQLite DB 저장 완료.", "sys")
    except Exception as e:
        conn.rollback()
        broadcaster.log(f"[CHAT_MOD] SQLite 저장 실패: {e}", "sys")
    finally:
        conn.close()

    return run_id, result

def run_chat_worker(req: ChatRequest):
    state.active_chat_running = True
    try:
        execute_chat_logic(req)
    except Exception as e:
        broadcaster.log(f"[CHAT_MOD] 백그라운드 에러: {e}", "sys")
    finally:
        state.active_chat_running = False

@router.post("/api/benchmark/chat")
def run_chat(req: ChatRequest, background_tasks: BackgroundTasks):
    if state.active_benchmark_running or state.active_chat_running:
        raise HTTPException(status_code=400, detail="A benchmark or chat session is already running.")
        
    background_tasks.add_task(run_chat_worker, req)
    return {"status": "started"}

@router.post("/api/benchmark/run-sync")
def run_benchmark_sync(req: SyncBenchmarkRequest):
    if state.active_benchmark_running or state.active_chat_running:
        raise HTTPException(status_code=400, detail="A benchmark or chat session is already running.")
    
    state.active_benchmark_running = True
    broadcaster.log(f"🚀 [Sync API] 벤치마크 가동 시퀀스 시작 (모드: {req.run_mode})", "bench")
    
    try:
        # 1. Smart SWAP 체크
        if state.last_booted_model != req.boot_config.model_name:
            broadcaster.log(f"🔄 [Sync API] 스왑 필요 감지: {state.last_booted_model} -> {req.boot_config.model_name}", "sys")
            config_dict = {
                "engine": req.boot_config.engine,
                "cpu_cores": req.boot_config.cpu_cores,
                "ram_mb": req.boot_config.ram_mb,
                "gpu_layers": req.boot_config.gpu_layers,
                "model_name": req.boot_config.model_name
            }
            state.engine.set_logger(lambda msg: broadcaster.log(f"{_ts()} {msg}", "sys"))
            success, msg = state.engine.boot_matrix(config_dict)
            if not success:
                state.boot_status = "ERROR"
                state.boot_message = f"스왑 실패: {msg}"
                broadcaster.log(f"❌ 스왑 실패: {msg}", "sys")
                raise HTTPException(status_code=500, detail=f"Engine swap failed: {msg}")
            state.boot_status = "ONLINE"
            state.boot_message = msg
            state.last_booted_model = req.boot_config.model_name

        # 2. 판정관 지정 동기화
        state.stress_config.judge_model = req.stress_config.judge_model
        state.save_config({"default_judge_model": req.stress_config.judge_model})

        # 3. 커스텀 태스크 목록 구성
        tasks_list = None
        if req.harness_tasks is not None:
            tasks_list = [t.dict() for t in req.harness_tasks]

        # 4. 실행 분기
        if "Stress" in req.run_mode or "Hard" in req.run_mode:
            run_id, results = _run_stress_mode(req)
        else:
            run_id, results = _run_inference_mode(req, tasks_list)

        # 5. 요약 생성
        scores = []
        tps_list = []
        ttft_list = []
        for r in results:
            try:
                scores.append(float(r["judge_score"]))
            except:
                pass
            if r.get("tps"):
                tps_list.append(r["tps"])
            if r.get("ttft_ms"):
                ttft_list.append(r["ttft_ms"])

        summary = {
            "total_tasks": len(results),
            "average_score": round(sum(scores)/len(scores), 2) if scores else 0.0,
            "average_tps": round(sum(tps_list)/len(tps_list), 2) if tps_list else 0.0,
            "average_ttft_ms": round(sum(ttft_list)/len(ttft_list), 1) if ttft_list else 0.0
        }

        return {
            "status": "success",
            "run_id": run_id,
            "summary": summary,
            "results": results
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        state.active_benchmark_running = False

@router.post("/api/benchmark/chat-sync")
def run_chat_sync(req: ChatRequest):
    if state.active_benchmark_running or state.active_chat_running:
        raise HTTPException(status_code=400, detail="A benchmark or chat session is already running.")
        
    state.active_chat_running = True
    try:
        run_id, result = execute_chat_logic(req)
        return {
            "status": "success",
            "run_id": run_id,
            "result": result
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        state.active_chat_running = False



# ─────────────────────────────────────────────────────────────────────────────
# EXPORT EXCEL & WORD
# ─────────────────────────────────────────────────────────────────────────────

@router.get("/api/reports/export/excel")
def export_excel(run_id: int = 0):
    conn = get_db_connection()
    if run_id > 0:
        query = "SELECT * FROM benchmark_results WHERE run_id = ?;"
        df = pd.read_sql_query(query, conn, params=(run_id,))
    else:
        # 전체 조인 내역
        query = """
        SELECT r.timestamp, r.model_name, r.engine_type, r.run_mode,
               s.task_name, s.category, s.prompt_text, s.response_text,
               s.ttft_ms, s.tps, s.avg_gpu_w, s.tokens_per_joule, s.judge_score, s.judge_reason
        FROM benchmark_runs r
        LEFT JOIN benchmark_results s ON r.id = s.run_id;
        """
        df = pd.read_sql_query(query, conn)
    conn.close()

    if df.empty:
        raise HTTPException(status_code=400, detail="No data found to export")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    df.to_excel(temp.name, index=False)
    temp.close()
    
    # 다운로드용 파일 전송
    return FileResponse(
        temp.name, 
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=f"AMEVA_Benchmark_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    )

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT EXPORT QUEUE SYSTEM
# ─────────────────────────────────────────────────────────────────────────────
import queue
import uuid

class DocExportTask:
    def __init__(self, task_id: str, run_id: int, model_name: str, use_llm_summary: bool):
        self.task_id = task_id
        self.run_id = run_id
        self.model_name = model_name
        self.use_llm_summary = use_llm_summary
        self.status = "pending"  # "pending", "processing", "completed", "failed", "cancelled"
        self.message = "대기 중"
        self.file_path = None
        self.filename = None
        self.error = None
        self.created_at = datetime.now().strftime("%H:%M:%S")

class DocExportQueueManager:
    def __init__(self):
        self.tasks = []
        self.lock = threading.Lock()
        self.task_queue = queue.Queue()
        self.worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
        self.worker_thread.start()

    def add_task(self, run_id: int, model_name: str, use_llm_summary: bool) -> str:
        task_id = f"doc_{uuid.uuid4().hex[:8]}"
        task = DocExportTask(task_id, run_id, model_name, use_llm_summary)
        with self.lock:
            self.tasks.append(task)
        self.task_queue.put(task_id)
        return task_id

    def cancel_task(self, task_id: str) -> bool:
        with self.lock:
            for task in self.tasks:
                if task.task_id == task_id:
                    if task.status in ("pending", "processing"):
                        task.status = "cancelled"
                        task.message = "취소됨"
                        return True
        return False

    def remove_task(self, task_id: str) -> bool:
        with self.lock:
            for i, task in enumerate(self.tasks):
                if task.task_id == task_id:
                    if task.file_path and os.path.exists(task.file_path):
                        try:
                            os.remove(task.file_path)
                        except:
                            pass
                    self.tasks.pop(i)
                    return True
        return False

    def get_task(self, task_id: str):
        with self.lock:
            for task in self.tasks:
                if task.task_id == task_id:
                    return task
        return None

    def get_all_tasks(self):
        with self.lock:
            return [
                {
                    "task_id": t.task_id,
                    "run_id": t.run_id,
                    "model_name": t.model_name,
                    "status": t.status,
                    "message": t.message,
                    "use_llm_summary": t.use_llm_summary,
                    "filename": t.filename,
                    "error": t.error,
                    "created_at": t.created_at
                }
                for t in self.tasks
            ]

    def _worker_loop(self):
        import time
        while True:
            try:
                task_id = self.task_queue.get()
                task = self.get_task(task_id)
                if not task:
                    self.task_queue.task_done()
                    continue
                if task.status == "cancelled":
                    self.task_queue.task_done()
                    continue

                with self.lock:
                    task.status = "processing"
                    task.message = "문서 정보 수집 중..."

                try:
                    self._generate_document(task)
                except Exception as e:
                    with self.lock:
                        task.status = "failed"
                        task.message = "문서 생성 실패"
                        task.error = str(e)
                        
                self.task_queue.task_done()
            except Exception as e:
                print(f"[DocExportQueueManager worker error] {e}")
                time.sleep(1.0)

    def _generate_document(self, task: DocExportTask):
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT * FROM benchmark_runs WHERE id = ?;", (task.run_id,))
        run_row = cursor.fetchone()
        if not run_row:
            conn.close()
            with self.lock:
                task.status = "failed"
                task.message = "해당 런을 찾을 수 없음"
            return
        run = dict(run_row)
        
        cursor.execute("SELECT * FROM benchmark_results WHERE run_id = ?;", (task.run_id,))
        results = [dict(r) for r in cursor.fetchall()]
        conn.close()

        if not results:
            with self.lock:
                task.status = "failed"
                task.message = "벤치마크 결과 없음"
            return

        if task.status == "cancelled":
            return

        summary_text = "요약을 비활성화했습니다."
        if task.use_llm_summary:
            judge_model = run.get("judge_model", state.config_data["default_judge_model"])
            with self.lock:
                task.message = "AI 요약 분석 중..."
            broadcaster.log(f"📝 보고서 분석용 요약 요청 시작: {judge_model}", "sys")
            
            prompt_data = (
                "아래 하네스 벤치마크 결과를 요약 분석하여, 해당 AI 엣지 디바이스 구동환경에서의 모델의 "
                "추론 속도, 지능 지표, 전력 효율성 측면의 강점과 한계점을 5줄 내외의 격식 있고 일목요연한 "
                "한글 요약글(Executive Summary)로 작성해 주세요. 문단 앞머리에 마크다운이나 특수 기호 없이 담백하게 작성하세요.\n\n"
                f"모델: {run['model_name']} ({run['engine_type']} 엔진)\n"
                f"설정: Cores={run['cpu_cores']}, RAM={run['ram_mb']}MB, Threads={run['threads']}, Context={run['n_ctx']}\n"
                "상세 태스크 지표:\n"
            )
            for r in results:
                prompt_data += f"- {r['task_name']} ({r['category']}): TTFT={r['ttft_ms']}ms, TPS={r['tps']}t/s, 전력={r['avg_gpu_w']}W, 점수={r['judge_score']}/10\n"

            try:
                from core.ollama_client import OllamaClient
                import json
                messages = [
                    {"role": "system", "content": "You are a professional IT technology reporting assistant. Answer in precise Korean."},
                    {"role": "user", "content": prompt_data}
                ]
                resp = OllamaClient.chat_stream(judge_model, messages, options={"temperature": 0.2})
                resp.raise_for_status()
                
                full_summary = ""
                for line in resp.iter_lines():
                    if task.status == "cancelled":
                        return
                    if line:
                        chunk = json.loads(line)
                        content = chunk.get("message", {}).get("content", "")
                        full_summary += content
                        with self.lock:
                            task.message = f"AI 요약 작성 중... ({len(full_summary)}자 생성)"
                summary_text = full_summary.strip()
                broadcaster.log("📝 보고서 요약 생성 완료.", "sys")
            except Exception as e:
                summary_text = f"요약 실패 (에러: {e})"
                broadcaster.log(f"❌ 보고서 요약 실패: {e}", "sys")

        if task.status == "cancelled":
            return

        with self.lock:
            task.message = "Word 문서 빌드 중..."

        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(11)

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("AMEVA AI 엣지 벤치마크 성능 평가 결과 보고서")
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 23, 42)

        doc.add_paragraph(f"보고서 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("1. 구동 환경 사양 및 모델 요약", level=2)
        meta_p = doc.add_paragraph()
        meta_p.add_run(f"• 대상 모델: {run['model_name']}\n").bold = True
        meta_p.add_run(f"• 추론 엔진: {run['engine_type']}\n")
        meta_p.add_run(f"• 할당 자원: CPU Cores={run['cpu_cores']} | RAM={run['ram_mb']}MB | GPU Layers={run['gpu_layers']}\n")
        meta_p.add_run(f"• 튜닝 파라미터: Threads={run['threads']} | n_ctx={run['n_ctx']} | Temp={run['temperature']}\n")
        meta_p.add_run(f"• 평가에 사용된 판정관: {run['judge_model']}")

        doc.add_heading("2. Executive Summary (AI 종합 분석 요약)", level=2)
        summary_box = doc.add_paragraph()
        summary_box.paragraph_format.left_indent = Inches(0.2)
        summary_box.paragraph_format.right_indent = Inches(0.2)
        run_sum = summary_box.add_run(summary_text)
        run_sum.italic = True
        run_sum.font.size = Pt(10.5)

        doc.add_heading("3. 태스크별 세부 측정 데이터", level=2)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Task ID'
        hdr_cells[1].text = 'TTFT'
        hdr_cells[2].text = 'TPS'
        hdr_cells[3].text = '평균전력'
        hdr_cells[4].text = '전성비(t/J)'
        hdr_cells[5].text = '판정점수'

        def set_cell_background(cell, color_hex):
            shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
            cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

        for cell in hdr_cells:
            set_cell_background(cell, "0F172A")
            for paragraph in cell.paragraphs:
                for run_cell in paragraph.runs:
                    run_cell.font.bold = True
                    run_cell.font.color.rgb = RGBColor(255, 255, 255)

        for item in results:
            if task.status == "cancelled":
                return
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['task_name'])
            row_cells[1].text = f"{item['ttft_ms']:.1f} ms"
            row_cells[2].text = f"{item['tps']:.2f} t/s"
            row_cells[3].text = f"{item['avg_gpu_w']:.1f} W"
            row_cells[4].text = f"{item['tokens_per_joule']:.3f} t/J"
            row_cells[5].text = str(item['judge_score'])

            try:
                score_num = float(item['judge_score'])
                if score_num >= 8.0:
                    set_cell_background(row_cells[5], "D1FAE5")
                elif score_num <= 4.0:
                    set_cell_background(row_cells[5], "FEE2E2")
            except:
                pass

        doc.add_heading("4. 상세 답변 및 판정 의견", level=2)
        for idx, item in enumerate(results):
            if task.status == "cancelled":
                return
            doc.add_heading(f"태스크 {idx+1}: {item['task_name']}", level=3)
            doc.add_paragraph(f"질문(Prompt): {item['prompt_text']}")
            
            resp_p = doc.add_paragraph()
            resp_p.add_run("답변(Response):\n").bold = True
            resp_p.add_run(item['response_text'])
            resp_p.style = 'Quote'
            
            reason_p = doc.add_paragraph()
            reason_p.add_run("판정 의견(Rationale):\n").bold = True
            reason_p.add_run(item['judge_reason'])

        if task.status == "cancelled":
            return

        import tempfile
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp.name)
        temp.close()

        with self.lock:
            task.status = "completed"
            task.message = "완료"
            task.file_path = temp.name
            sanitized_model = "".join([c if c.isalnum() or c in ("-", "_") else "_" for c in run['model_name']])
            task.filename = f"AMEVA_Performance_Report_{sanitized_model}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

doc_queue_manager = DocExportQueueManager()

@router.post("/api/reports/export/word")
def export_word(req: ExportWordRequest):
    # Keep synchronous implementation for compatibility
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM benchmark_runs WHERE id = ?;", (req.run_id,))
    run_row = cursor.fetchone()
    if not run_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Run not found")
    run = dict(run_row)
    cursor.execute("SELECT * FROM benchmark_results WHERE run_id = ?;", (req.run_id,))
    results = [dict(r) for r in cursor.fetchall()]
    conn.close()

    if not results:
        raise HTTPException(status_code=400, detail="No results found")

    summary_text = "요약을 비활성화했습니다."
    if req.use_llm_summary:
        judge_model = run.get("judge_model", state.config_data["default_judge_model"])
        broadcaster.log(f"📝 보고서 분석용 요약 요청 시작: {judge_model}", "sys")
        prompt_data = (
            "아래 하네스 벤치마크 결과를 요약 분석하여, 해당 AI 엣지 디바이스 구동환경에서의 모델의 "
            "추론 속도, 지능 지표, 전력 효율성 측면의 강점과 한계점을 5줄 내외의 격식 있고 일목요연한 "
            "한글 요약글(Executive Summary)로 작성해 주세요. 문단 앞머리에 마크다운이나 특수 기호 없이 담백하게 작성하세요.\n\n"
            f"모델: {run['model_name']} ({run['engine_type']} 엔진)\n"
            f"설정: Cores={run['cpu_cores']}, RAM={run['ram_mb']}MB, Threads={run['threads']}, Context={run['n_ctx']}\n"
            "상세 태스크 지표:\n"
        )
        for r in results:
            prompt_data += f"- {r['task_name']} ({r['category']}): TTFT={r['ttft_ms']}ms, TPS={r['tps']}t/s, 전력={r['avg_gpu_w']}W, 점수={r['judge_score']}/10\n"

        try:
            from core.ollama_client import OllamaClient
            messages = [
                {"role": "system", "content": "You are a professional IT technology reporting assistant. Answer in precise Korean."},
                {"role": "user", "content": prompt_data}
            ]
            resp = OllamaClient.chat_stream(judge_model, messages, options={"temperature": 0.2})
            resp.raise_for_status()
            full_summary = ""
            for line in resp.iter_lines():
                if line:
                    chunk = json.loads(line)
                    content = chunk.get("message", {}).get("content", "")
                    full_summary += content
            summary_text = full_summary.strip()
            broadcaster.log("📝 보고서 요약 생성 완료.", "sys")
        except Exception as e:
            summary_text = f"요약 실패 (에러: {e})"
            broadcaster.log(f"❌ 보고서 요약 실패: {e}", "sys")

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("AMEVA AI 엣지 벤치마크 성능 평가 결과 보고서")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(f"보고서 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_heading("1. 구동 환경 사양 및 모델 요약", level=2)
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"• 대상 모델: {run['model_name']}\n").bold = True
    meta_p.add_run(f"• 추론 엔진: {run['engine_type']}\n")
    meta_p.add_run(f"• 할당 자원: CPU Cores={run['cpu_cores']} | RAM={run['ram_mb']}MB | GPU Layers={run['gpu_layers']}\n")
    meta_p.add_run(f"• 튜닝 파라미터: Threads={run['threads']} | n_ctx={run['n_ctx']} | Temp={run['temperature']}\n")
    meta_p.add_run(f"• 평가에 사용된 판정관: {run['judge_model']}")

    doc.add_heading("2. Executive Summary (AI 종합 분석 요약)", level=2)
    summary_box = doc.add_paragraph()
    summary_box.paragraph_format.left_indent = Inches(0.2)
    summary_box.paragraph_format.right_indent = Inches(0.2)
    run_sum = summary_box.add_run(summary_text)
    run_sum.italic = True
    run_sum.font.size = Pt(10.5)

    doc.add_heading("3. 태스크별 세부 측정 데이터", level=2)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Task ID'
    hdr_cells[1].text = 'TTFT'
    hdr_cells[2].text = 'TPS'
    hdr_cells[3].text = '평균전력'
    hdr_cells[4].text = '전성비(t/J)'
    hdr_cells[5].text = '판정점수'

    def set_cell_background(cell, color_hex):
        shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

    for cell in hdr_cells:
        set_cell_background(cell, "0F172A")
        for paragraph in cell.paragraphs:
            for run_cell in paragraph.runs:
                run_cell.font.bold = True
                run_cell.font.color.rgb = RGBColor(255, 255, 255)

    for item in results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['task_name'])
        row_cells[1].text = f"{item['ttft_ms']:.1f} ms"
        row_cells[2].text = f"{item['tps']:.2f} t/s"
        row_cells[3].text = f"{item['avg_gpu_w']:.1f} W"
        row_cells[4].text = f"{item['tokens_per_joule']:.3f} t/J"
        row_cells[5].text = str(item['judge_score'])

        try:
            score_num = float(item['judge_score'])
            if score_num >= 8.0:
                set_cell_background(row_cells[5], "D1FAE5")
            elif score_num <= 4.0:
                set_cell_background(row_cells[5], "FEE2E2")
        except:
            pass

    doc.add_heading("4. 상세 답변 및 판정 의견", level=2)
    for idx, item in enumerate(results):
        doc.add_heading(f"태스크 {idx+1}: {item['task_name']}", level=3)
        doc.add_paragraph(f"질문(Prompt): {item['prompt_text']}")
        resp_p = doc.add_paragraph()
        resp_p.add_run("답변(Response):\n").bold = True
        resp_p.add_run(item['response_text'])
        resp_p.style = 'Quote'
        reason_p = doc.add_paragraph()
        reason_p.add_run("판정 의견(Rationale):\n").bold = True
        reason_p.add_run(item['judge_reason'])

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()

    return FileResponse(
        temp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"AMEVA_Performance_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

@router.post("/api/reports/export/word/queue")
def export_word_queue(req: ExportWordRequest):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT model_name FROM benchmark_runs WHERE id = ?;", (req.run_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Run not found")
    model_name = row["model_name"]
    
    task_id = doc_queue_manager.add_task(req.run_id, model_name, req.use_llm_summary)
    return {"task_id": task_id, "status": "pending"}

@router.get("/api/reports/export/word/queue")
def get_word_queue():
    return doc_queue_manager.get_all_tasks()

@router.delete("/api/reports/export/word/queue/{task_id}")
def cancel_or_remove_word_task(task_id: str):
    cancelled = doc_queue_manager.cancel_task(task_id)
    removed = doc_queue_manager.remove_task(task_id)
    return {"task_id": task_id, "cancelled": cancelled, "removed": removed}

@router.get("/api/reports/export/word/download/{task_id}")
def download_word_file(task_id: str):
    task = doc_queue_manager.get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task.status != "completed":
        raise HTTPException(status_code=400, detail=f"Task is in status '{task.status}', not ready for download")
    if not task.file_path or not os.path.exists(task.file_path):
        raise HTTPException(status_code=404, detail="Generated file not found on disk")
        
    return FileResponse(
        task.file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=task.filename
    )
