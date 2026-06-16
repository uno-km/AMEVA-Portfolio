import os
import json
import time
import traceback
import threading
from datetime import datetime

from core.llm_engine import LlamaInferenceCore
from core.sre import logger
from core.config import MEMORY_DIR
from core.security import enforce_sandbox
from core.parser import StrictParser
from agents.schemas import (
    PM_SCHEMA, ARCHITECT_SCHEMA,
    ARCHITECT_SPEC_SCHEMA, DEV_SINGLE_FILE_SCHEMA,
    REVIEW_SCHEMA, WORKER_SCHEMA, PROMPTS
)
from core.database import DatabaseManager


class AgentWorker(threading.Thread):
    """에이전트 하나의 실행 단위. 역할에 따라 내부 처리 흐름이 분기된다."""

    def __init__(self, agent_id, role_prompt, task_data, on_done=None, on_fail=None, on_stream=None):
        super().__init__(daemon=True)
        self.agent_id    = agent_id
        self.role_prompt = role_prompt
        self.task_data   = task_data
        self.on_done     = on_done
        self.on_fail     = on_fail
        self.on_stream   = on_stream
        self.heartbeat   = time.time()
        self.llm_core    = LlamaInferenceCore.get_instance()
        self._stop_event = threading.Event()

    # ── 중단 제어 ──────────────────────────────────────────
    def requestInterruption(self):
        self._stop_event.set()

    def isInterruptionRequested(self):
        return self._stop_event.is_set()

    # ── 스트리밍 콜백 ───────────────────────────────────────
    def _stream_cb(self, delta):
        if self.isInterruptionRequested():
            raise Exception("Agent worker interrupted by user request.")
        if self.on_stream:
            self.on_stream(self.agent_id, delta)

    # ── LLM 호출 공통 래퍼 ─────────────────────────────────
    def _call_llm(self, full_prompt, schema):
        return self.llm_core.generate(self.role_prompt, full_prompt, schema, stream_callback=self._stream_cb)

    # ── DB 컨텍스트 준비 ────────────────────────────────────
    def _prepare_context(self):
        workflow_id  = self.task_data.get("workflow_id")
        original_goal = self.task_data.get("original_goal", "목표 미지정")
        instruction  = self.task_data.get("instruction", "주어진 목표를 완수하십시오.")
        sender_id    = self.task_data.get("sender_id", "unknown")

        engine = LlamaInferenceCore.get_instance()
        current_model = os.path.basename(engine.current_model_path) if engine.current_model_path else "Unknown"

        task_seq_id = 0
        past_mem    = "이전 히스토리 없음."
        if workflow_id:
            task_seq_id = DatabaseManager.create_task(
                workflow_id, self.agent_id, instruction, current_model, sender_id
            )
            past_mem = DatabaseManager.get_workflow_context(workflow_id)

        return workflow_id, original_goal, instruction, task_seq_id, past_mem

    # ── 메인 실행 진입점 ────────────────────────────────────
    def run(self):
        try:
            if self.isInterruptionRequested():
                return
            self.heartbeat = time.time()

            # 에이전트 역할별 분기
            if self.agent_id == "pm":
                self._run_pm()
            elif self.agent_id == "architect":
                self._run_architect()
            elif self.agent_id == "dev":
                self._run_dev()
            else:
                # tester, secretary — 기존 단순 워커
                self._run_generic()

        except Exception as e:
            if "interrupted by user request" in str(e).lower():
                logger.info(f"WORKER [{self.agent_id}]: Interrupted by user. Exiting cleanly.")
                if self.on_fail:
                    self.on_fail(self.agent_id, "작업 중단됨")
                return
            tb = traceback.format_exc()
            logger.error(f"WORKER FATAL [{self.agent_id}]: {tb}")
            wid = self.task_data.get("workflow_id", "")
            if wid:
                DatabaseManager.log_exception(wid, 0, self.agent_id, str(e), tb)
            if self.on_fail:
                self.on_fail(self.agent_id, f"워커 치명적 오류: {str(e)}")

    # ══════════════════════════════════════════════════════════
    # PM 에이전트
    # ══════════════════════════════════════════════════════════
    def _run_pm(self):
        workflow_id, original_goal, instruction, task_seq_id, past_mem = self._prepare_context()

        full_prompt = (
            f"### [PAST LOGS & CONTEXT]\n{past_mem}\n\n"
            f"### [CURRENT MISSION]\n{instruction}"
        )

        result_json, usage = self._call_llm(full_prompt, PM_SCHEMA)
        if self.isInterruptionRequested():
            return

        # 파싱 실패 폴백
        if result_json.get("status") == 500:
            result_json = {
                "overall_plan": "파싱 오류. 아키텍트부터 재시작합니다.",
                "thought":      "JSON 파싱 오류. 기본 폴백으로 복구합니다.",
                "summary":      "파싱 실패 폴백",
                "next_action":  {"target": "architect", "instruction": instruction}
            }

        # PM 유효성 검사
        try:
            self._validate_pm_plan(result_json)
        except ValueError as ve:
            logger.warning(f"PM validation failed: {ve}. Fallback.")
            result_json["next_action"] = {"target": "architect", "instruction": instruction}

        next_action = result_json.get("next_action", {})
        target      = next_action.get("target", "secretary")
        next_instr  = next_action.get("instruction", "최종 보고를 수행하십시오.")

        result_json["message"] = f"PM 결정: {target.upper()} 에게 위임 ({result_json.get('summary', '')})"

        # 다음 task_data 구성
        next_task = {
            "target":       target,
            "instruction":  next_instr,
            "plan":         [],
            "workflow_id":  workflow_id,
            "original_goal": original_goal,
        }

        # Architect / Dev / Tester → 완료 후 PM에게 복귀 예약
        if target not in ("secretary",):
            next_task["plan"] = [{
                "target": "pm",
                "instruction": (
                    f"Goal: {original_goal}. "
                    f"이전 에이전트({target})의 산출물과 피드백을 분석하여 다음 재귀 추론 단계를 결정하십시오."
                )
            }]

        # Dev에게 위임할 때 아키텍트의 spec을 함께 전달
        if target == "dev":
            next_task["architect_spec"] = self.task_data.get("architect_spec", {})

        if workflow_id:
            DatabaseManager.log_task_dtl(workflow_id, task_seq_id, self.agent_id, "execution", result_json, result_json.get("message", ""))

        next_task["hop_count"]       = self.task_data.get("hop_count", 0) + 1
        next_task["visited_targets"] = list(self.task_data.get("visited_targets", [])) + [self.agent_id]

        if self.on_done:
            self.on_done(self.agent_id, result_json, next_task, usage)

    # ══════════════════════════════════════════════════════════
    # Architect 에이전트 (설계 or 코드 리뷰 모드)
    # ══════════════════════════════════════════════════════════
    def _run_architect(self):
        workflow_id, original_goal, instruction, task_seq_id, past_mem = self._prepare_context()

        review_mode = self.task_data.get("review_mode", False)

        if review_mode:
            # ── 코드 리뷰 모드 ──────────────────────────────
            code_structure = DatabaseManager.get_code_structure(workflow_id)

            # 파일 내용을 프롬프트에 포함 (파일별로 내용 주입)
            code_review_prompt = "아래 생성된 파일들을 하나씩 검토하고 REVIEW_SCHEMA 형식으로 verdict와 issues를 출력하십시오.\n\n"
            for f in code_structure:
                content = DatabaseManager.get_code_file_content(workflow_id, f["file_path"])
                # 너무 길면 잘라냄
                snippet = content[:1500] + "\n... (truncated)" if len(content) > 1500 else content
                code_review_prompt += f"### [{f['file_path']}]\n```\n{snippet}\n```\n\n"

            full_prompt = (
                f"### [PAST LOGS & CONTEXT]\n{past_mem}\n\n"
                f"### [CODE REVIEW REQUEST]\n{code_review_prompt}"
            )
            result_json, usage = self._call_llm(full_prompt, REVIEW_SCHEMA)

            if result_json.get("status") == 500:
                result_json = {"verdict": "rejected", "issues": [{"file_path": "unknown", "description": "LLM 파싱 오류"}], "message": "파싱 실패로 거부"}

            result_json["message"] = f"코드 리뷰 완료: {result_json.get('verdict', 'unknown').upper()}"
            result_json["agent_mode"] = "code_review"

        else:
            # ── 설계 문서 생성 모드 ─────────────────────────
            spec_prompt = (
                "PM의 지시를 바탕으로 ARCHITECT_SPEC_SCHEMA 형식의 개발요구문서(JSON Design Spec)를 작성하십시오.\n"
                "코드는 절대 포함하지 마십시오. 파일 구조, 역할, 설명, 의존성만 정의하십시오.\n"
                "각 파일은 단일 책임 원칙을 지켜 분리하십시오.\n\n"
                f"PM 지시사항:\n{instruction}"
            )

            full_prompt = (
                f"### [PAST LOGS & CONTEXT]\n{past_mem}\n\n"
                f"### [DESIGN SPEC REQUEST]\n{spec_prompt}"
            )
            result_json, usage = self._call_llm(full_prompt, ARCHITECT_SPEC_SCHEMA)

            if result_json.get("status") == 500:
                result_json = {
                    "project_root": "CodeGod_Workspace/project",
                    "requirements": [instruction],
                    "directory_tree": ["main.py"],
                    "file_specs": [{"file_path": "main.py", "concept": "메인 진입점", "description": instruction, "dependencies": [], "exports": []}],
                    "message": "파싱 실패 폴백 스펙"
                }

            result_json["message"] = f"설계 완료: {len(result_json.get('file_specs', []))}개 파일 명세 작성"
            result_json["agent_mode"] = "design_spec"

        if workflow_id:
            DatabaseManager.log_task_dtl(workflow_id, task_seq_id, self.agent_id, "execution", result_json, result_json.get("message", ""))

        next_task = self._build_next_task_from_plan(workflow_id, original_goal)

        # Architect가 spec을 완성했으면 next_task에 spec 첨부 (PM이 Dev에게 전달할 때 사용)
        if not review_mode and "file_specs" in result_json:
            if next_task:
                next_task["architect_spec"] = result_json
            # PM에 복귀하는 plan에도 spec 첨부
            for p in (next_task or {}).get("plan", []):
                if p.get("target") == "pm":
                    p["architect_spec"] = result_json

        if self.on_done:
            self.on_done(self.agent_id, result_json, next_task, usage)

    # ══════════════════════════════════════════════════════════
    # Dev 에이전트 — 파일 하나씩 재귀 루프
    # ══════════════════════════════════════════════════════════
    def _run_dev(self):
        workflow_id, original_goal, instruction, task_seq_id, past_mem = self._prepare_context()

        # ── 수정 모드 (리뷰 이슈가 있는 경우) ───────────────
        review_issues = self.task_data.get("code_review_issues", [])
        if review_issues:
            self._run_dev_fix(workflow_id, original_goal, instruction, task_seq_id, past_mem, review_issues)
            return

        # ── 신규 구현 모드 ───────────────────────────────────
        spec         = self.task_data.get("architect_spec", {})
        file_specs   = list(self.task_data.get("remaining_file_specs", spec.get("file_specs", [])))
        project_root = spec.get("project_root", "CodeGod_Workspace")
        review_round = self.task_data.get("review_round", 0)

        if not file_specs:
            # 모든 파일 완료 → 구조 보고서 작성 후 반환
            self._dev_submit_report(workflow_id, original_goal, task_seq_id, project_root, review_round, past_mem)
            return

        # 현재 파일 처리
        current_spec   = file_specs[0]
        remaining_specs = file_specs[1:]
        current_file_path = current_spec.get("file_path", "unknown.py")

        retry_count = 0
        max_retries = 3
        security_warning = ""

        while retry_count < max_retries:
            dev_prompt = (
                f"### [PAST LOGS & CONTEXT]\n{past_mem}\n\n"
                f"### [DEV TASK — 단일 파일 구현]\n"
                f"Project Root: {project_root}\n"
                f"구현할 파일: {current_file_path}\n\n"
                f"파일 명세:\n{json.dumps(current_spec, ensure_ascii=False, indent=2)}\n\n"
                f"전체 프로젝트 요구사항:\n{json.dumps(spec.get('requirements', []), ensure_ascii=False)}\n\n"
            )
            if security_warning:
                dev_prompt += (
                    f"\n[⚠️ 보안 검사 경고 - 보안 위반 감지]\n"
                    f"이전 코드 작성 시도 중 보안 규칙 위반이 감지되어 시스템에 의해 차단되었습니다:\n"
                    f"-> {security_warning}\n"
                    f"이 문제를 즉시 해결하십시오. 위험 코드(예: eval(), exec(), os.system(), subprocess 등)를 절대 사용하지 않고 안전하게 대체 로직을 구현해야 합니다.\n\n"
                )
            
            dev_prompt += (
                f"[CRITICAL] DEV_SINGLE_FILE_SCHEMA (file_path, content, message) 형식으로만 출력하십시오.\n"
                f"[CRITICAL] JSON 내 Python 코드에서 큰따옴표는 반드시 \\\"로 이스케이프하십시오."
            )

            result_json, usage = self._call_llm(dev_prompt, DEV_SINGLE_FILE_SCHEMA)

            if result_json.get("status") == 500:
                result_json = {
                    "file_path": current_file_path,
                    "content":   f"# TODO: {current_spec.get('description', 'Implement this file')}\npass",
                    "message":   f"LLM 파싱 실패 — 폴백 스텁 생성: {current_file_path}"
                }

            # 물리 파일 저장 시도
            try:
                saved_path = self._save_dev_file(project_root, result_json, current_file_path)
                break  # 성공 시 루프 탈출
            except ValueError as ve:
                if "보안 위험" in str(ve) or "security" in str(ve).lower() or "blocked" in str(ve).lower() or "MALICIOUS" in str(ve):
                    retry_count += 1
                    security_warning = str(ve)
                    logger.warning(f"DEV: [{current_file_path}] 보안 차단 발생 (시도 {retry_count}/{max_retries}): {ve}. 재시도 프롬프트 전송.")
                    if retry_count >= max_retries:
                        raise ValueError(f"보안 위험 요소 감지 및 자율 수정 {max_retries}회 실패로 중단되었습니다: {ve}")
                    time.sleep(1)
                else:
                    raise ve

        # DB 저장 (code_files)
        if workflow_id:
            DatabaseManager.upsert_code_file(
                workflow_id, task_seq_id,
                result_json.get("file_path", current_file_path),
                result_json.get("content", ""),
                review_round
            )
            DatabaseManager.log_task_dtl(workflow_id, task_seq_id, self.agent_id, "file_created",
                                          result_json, result_json.get("message", ""))

        logger.info(f"DEV: [{current_file_path}] 완료. 남은 파일: {len(remaining_specs)}개")

        # ── 재귀: 남은 파일이 있으면 자신에게 다시 dispatch ─
        if remaining_specs:
            next_task = {
                "target":               "dev",
                "instruction":          instruction,
                "workflow_id":          workflow_id,
                "original_goal":        original_goal,
                "architect_spec":       spec,
                "remaining_file_specs": remaining_specs,
                "review_round":         review_round,
                "plan":                 self.task_data.get("plan", []),
                "hop_count":            self.task_data.get("hop_count", 0) + 1,
                "visited_targets":      list(self.task_data.get("visited_targets", [])),
                "sender_id":            "dev"
            }
        else:
            # ── 모든 파일 완료 → 보고서 생성 후 PM 복귀 ────
            next_task = self._dev_build_completion_task(workflow_id, original_goal, project_root)

        if self.on_done:
            self.on_done(self.agent_id, result_json, next_task, usage)

    def _run_dev_fix(self, workflow_id, original_goal, instruction, task_seq_id, past_mem, review_issues):
        """리뷰 이슈를 받아 해당 파일만 수정하는 모드"""
        spec         = self.task_data.get("architect_spec", {})
        project_root = spec.get("project_root", "CodeGod_Workspace")
        review_round = self.task_data.get("review_round", 1)

        # 이슈 목록 중 첫 번째 파일부터 수정
        current_issue = review_issues[0]
        remaining_issues = review_issues[1:]

        file_path    = current_issue.get("file_path", "")
        existing_code = DatabaseManager.get_code_file_content(workflow_id, file_path) if workflow_id else ""

        retry_count = 0
        max_retries = 3
        security_warning = ""

        while retry_count < max_retries:
            fix_prompt = (
                f"### [PAST LOGS & CONTEXT]\n{past_mem}\n\n"
                f"### [DEV TASK — 코드 수정]\n"
                f"수정할 파일: {file_path}\n\n"
                f"리뷰 이슈:\n{json.dumps(current_issue, ensure_ascii=False, indent=2)}\n\n"
                f"현재 코드:\n```python\n{existing_code}\n```\n\n"
            )
            if security_warning:
                fix_prompt += (
                    f"\n[⚠️ 보안 검사 경고 - 보안 위반 감지]\n"
                    f"이전 코드 작성 시도 중 보안 규칙 위반이 감지되어 시스템에 의해 차단되었습니다:\n"
                    f"-> {security_warning}\n"
                    f"이 문제를 즉시 해결하십시오. 위험 코드(예: eval(), exec(), os.system(), subprocess 등)를 절대 사용하지 않고 안전하게 대체 로직을 구현해야 합니다.\n\n"
                )
            fix_prompt += (
                f"[CRITICAL] 수정된 전체 코드를 DEV_SINGLE_FILE_SCHEMA 형식으로 출력하십시오.\n"
                f"[CRITICAL] JSON 내 Python 코드의 큰따옴표는 반드시 \\\"로 이스케이프하십시오."
            )

            result_json, usage = self._call_llm(fix_prompt, DEV_SINGLE_FILE_SCHEMA)

            if result_json.get("status") == 500:
                result_json = {
                    "file_path": file_path,
                    "content":   existing_code,
                    "message":   f"수정 파싱 실패 — 원본 유지: {file_path}"
                }

            # 물리 파일 덮어쓰기 시도
            try:
                self._save_dev_file(project_root, result_json, file_path)
                break
            except ValueError as ve:
                if "보안 위험" in str(ve) or "security" in str(ve).lower() or "blocked" in str(ve).lower() or "MALICIOUS" in str(ve):
                    retry_count += 1
                    security_warning = str(ve)
                    logger.warning(f"DEV FIX: [{file_path}] 보안 차단 발생 (시도 {retry_count}/{max_retries}): {ve}. 재시도 프롬프트 전송.")
                    if retry_count >= max_retries:
                        raise ValueError(f"보안 위험 요소 감지 및 자율 수정 {max_retries}회 실패로 중단되었습니다: {ve}")
                    time.sleep(1)
                else:
                    raise ve

        # DB 업데이트
        if workflow_id:
            DatabaseManager.upsert_code_file(
                workflow_id, task_seq_id,
                result_json.get("file_path", file_path),
                result_json.get("content", ""),
                review_round
            )
            DatabaseManager.log_task_dtl(workflow_id, task_seq_id, self.agent_id, "file_fixed",
                                          result_json, result_json.get("message", ""))

        # 남은 이슈가 있으면 자신에게 재귀
        if remaining_issues:
            next_task = {
                "target":              "dev",
                "instruction":         instruction,
                "workflow_id":         workflow_id,
                "original_goal":       original_goal,
                "architect_spec":      spec,
                "code_review_issues":  remaining_issues,
                "review_round":        review_round,
                "plan":                self.task_data.get("plan", []),
                "hop_count":           self.task_data.get("hop_count", 0) + 1,
                "visited_targets":     list(self.task_data.get("visited_targets", [])),
                "sender_id":           "dev"
            }
        else:
            # 수정 완료 → 다시 리뷰 요청 (PM 복귀)
            next_task = self._dev_build_completion_task(workflow_id, original_goal, project_root)

        if self.on_done:
            self.on_done(self.agent_id, result_json, next_task, usage)

    def _save_dev_file(self, project_root, result_json, fallback_path):
        """물리 파일 디스크 저장"""
        rel_path = result_json.get("file_path", fallback_path) or fallback_path

        # LLM이 확장자 없이 디렉토리 경로만 반환한 경우 방어
        if not os.path.splitext(rel_path)[1]:
            # fallback_path에 확장자가 있으면 그걸 사용, 없으면 .py 붙임
            if os.path.splitext(fallback_path)[1]:
                rel_path = fallback_path
            else:
                rel_path = rel_path.rstrip("/\\") + ".py"
            logger.warning(f"DEV: 확장자 없는 file_path 감지 → 폴백 사용: {rel_path}")

        full_path = os.path.join(project_root, rel_path) if not os.path.isabs(rel_path) else rel_path
        safe_path = enforce_sandbox(full_path)
        os.makedirs(os.path.dirname(safe_path), exist_ok=True)
        content = result_json.get("content", "")
        content = StrictParser.sanitize_code(content, rel_path, "dev")
        with open(safe_path, "w", encoding="utf-8") as f:
            f.write(content)
        return safe_path

    def _dev_build_completion_task(self, workflow_id, original_goal, project_root):
        """Dev가 모든 파일 완료 후 PM에 복귀하는 task 구성 (코드 구조 보고서 포함)"""
        code_structure = DatabaseManager.get_code_structure(workflow_id) if workflow_id else []
        structure_json = json.dumps(code_structure, ensure_ascii=False)

        return {
            "target": "pm",
            "instruction": (
                f"Goal: {original_goal}. "
                f"개발자(Dev)가 모든 파일 구현을 완료했습니다. 아키텍트에게 코드 리뷰를 지시하십시오. "
                f"구현된 코드 구조 보고서:\n{structure_json}"
            ),
            "workflow_id":   workflow_id,
            "original_goal": original_goal,
            "plan":          [],
            "hop_count":     self.task_data.get("hop_count", 0) + 1,
            "visited_targets": list(self.task_data.get("visited_targets", [])) + ["dev"],
            "sender_id":     "dev"
        }

    def _dev_submit_report(self, workflow_id, original_goal, task_seq_id, project_root, review_round, past_mem):
        """file_specs가 없어서 즉시 보고하는 경우 (빈 spec 폴백)"""
        result_json = {"file_path": "report.md", "content": "No files to implement.", "message": "빈 스펙 수신"}
        next_task = self._dev_build_completion_task(workflow_id, original_goal, project_root)
        if self.on_done:
            self.on_done(self.agent_id, result_json, next_task, {})

    # ══════════════════════════════════════════════════════════
    # Generic 에이전트 (tester, secretary)
    # ══════════════════════════════════════════════════════════
    def _run_generic(self):
        workflow_id, original_goal, instruction, task_seq_id, past_mem = self._prepare_context()

        full_prompt = (
            f"### [PAST LOGS & CONTEXT]\n{past_mem}\n\n"
            f"### [CURRENT MISSION]\n{instruction}"
        )
        result_json, usage = self._call_llm(full_prompt, WORKER_SCHEMA)

        if result_json.get("status") == 500:
            result_json = {"status": 200, "file_name": "output.md", "content": "LLM 파싱 오류", "message": "파싱 실패 폴백"}

        # ── Tester 샌드박스 테스트 실행 ──
        if self.agent_id == "tester" and result_json.get("status") == 200 and "file_name" in result_json:
            import subprocess
            import shutil
            import venv
            env_dir = enforce_sandbox(".test_venv")
            try:
                safe_path = enforce_sandbox(result_json["file_name"])
                os.makedirs(os.path.dirname(safe_path), exist_ok=True)
                with open(safe_path, "w", encoding="utf-8") as f:
                    f.write(result_json.get("content", ""))
                
                logger.info("TESTER: Creating sandbox venv...")
                venv.create(env_dir, with_pip=False)
                
                if os.name == 'nt':
                    py_exe = os.path.join(env_dir, "Scripts", "python.exe")
                else:
                    py_exe = os.path.join(env_dir, "bin", "python")
                    
                logger.info(f"TESTER: Running {safe_path} in sandbox...")
                proc = subprocess.run([py_exe, safe_path], capture_output=True, text=True, timeout=15)
                
                if proc.returncode != 0:
                    result_json["message"] = f"[TEST FAILED] {proc.stderr[:1000]}"
                else:
                    result_json["message"] = f"테스트 성공: {proc.stdout.strip()[:200]}"
            except subprocess.TimeoutExpired:
                result_json["message"] = "[TEST FAILED] Timeout Exceeded (15s)."
            except Exception as e:
                result_json["message"] = f"[TEST ERROR] {e}"
            finally:
                logger.info("TESTER: Cleaning up sandbox venv...")
                shutil.rmtree(env_dir, ignore_errors=True)

        # ── Secretary Word 보고서 생성 ──
        if self.agent_id == "secretary":
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.oxml.ns import qn
                
                doc = Document()
                # 폰트 깨짐 방지: 맑은 고딕 설정
                style = doc.styles['Normal']
                style.font.name = '맑은 고딕'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                doc.add_heading('프로젝트 최종 보고서', 0)
                doc.add_heading('1. 개요 및 요약', level=1)
                doc.add_paragraph(result_json.get("content", ""))
                
                # 생성된 파일 목록 스캔
                doc.add_heading('2. 생성된 파일 목록', level=1)
                workspace_dir = enforce_sandbox("")
                file_count = 0
                for root, dirs, files in os.walk(workspace_dir):
                    if ".test_venv" in root or "__pycache__" in root: continue
                    for f in files:
                        if f.endswith(".docx") or f.endswith(".log"): continue
                        doc.add_paragraph(f"- {os.path.relpath(os.path.join(root, f), workspace_dir)}")
                        file_count += 1
                if file_count == 0:
                    doc.add_paragraph("생성된 파일이 없습니다.")
                
                # 워크플로 대화 기록 추가
                doc.add_heading('3. 상세 에이전트 대화 및 작업 로그', level=1)
                history = DatabaseManager.get_workflow_context(workflow_id) if workflow_id else "로그 없음"
                doc.add_paragraph(history)
                
                report_name = f"Final_Report_{datetime.now().strftime('%H%M%S')}.docx"
                report_path = enforce_sandbox(report_name)
                doc.save(report_path)
                result_json["message"] = f"최종 보고서(.docx)가 성공적으로 생성되었습니다!\n파일명: {report_name}\n"
            except Exception as e:
                logger.error(f"SECRETARY DOCX ERROR: {e}")
                result_json["message"] = f"보고서 생성 실패: {e}"

        if workflow_id:
            DatabaseManager.log_task_dtl(workflow_id, task_seq_id, self.agent_id, "execution",
                                          result_json, result_json.get("message", "완료"))

        # Tester 실패 시 PM에게 반려 피드백 루프
        if self.agent_id == "tester" and "[TEST FAILED]" in result_json.get("message", ""):
            next_task = {
                "target": "pm",
                "instruction": f"테스트 실패: {result_json['message']}\n에러 로그를 분석하여 개발자에게 코드 수정을 지시하십시오.",
                "workflow_id": workflow_id,
                "original_goal": original_goal,
                "plan": [],
                "hop_count": self.task_data.get("hop_count", 0) + 1,
                "visited_targets": list(self.task_data.get("visited_targets", [])) + ["tester"],
                "sender_id": "tester"
            }
        else:
            next_task = self._build_next_task_from_plan(workflow_id, original_goal)

        if self.on_done:
            self.on_done(self.agent_id, result_json, next_task, usage)

    # ══════════════════════════════════════════════════════════
    # 헬퍼
    # ══════════════════════════════════════════════════════════
    def _build_next_task_from_plan(self, workflow_id, original_goal):
        """task_data 내 plan 배열에서 다음 항목을 꺼내 next_task를 구성한다."""
        next_plan = list(self.task_data.get("plan", []))
        if not next_plan:
            return None
        next_task = next_plan.pop(0)
        next_task["plan"]            = next_plan
        next_task["workflow_id"]     = workflow_id
        next_task["original_goal"]   = original_goal
        next_task["hop_count"]       = self.task_data.get("hop_count", 0) + 1
        next_task["visited_targets"] = list(self.task_data.get("visited_targets", [])) + [self.agent_id]
        # architect_spec 승계
        if "architect_spec" in self.task_data:
            next_task.setdefault("architect_spec", self.task_data["architect_spec"])
        return next_task

    def _validate_pm_plan(self, result_json):
        """PM 결과물 구조 검증"""
        if "next_action" not in result_json or not isinstance(result_json["next_action"], dict):
            raise ValueError("PM agent returned invalid or missing next_action structure.")
        action  = result_json["next_action"]
        target  = action.get("target")
        instr   = action.get("instruction")
        valid_targets = {"architect", "dev", "tester", "secretary"}
        if target not in valid_targets:
            raise ValueError(f"Invalid next_action target '{target}'. Allowed: {valid_targets}")
        if not instr or not isinstance(instr, str):
            raise ValueError("next_action instruction must be a non-empty string.")
