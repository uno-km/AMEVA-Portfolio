import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_portfolio():
    doc_path = r"c:\ameva\AMEVA-Portfolio\combined_profile_summary.docx"
    
    # Create a new document to build from scratch for perfect layout control
    doc = docx.Document()
    
    # Configure default style (Malgun Gothic)
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black for premium look
    
    # Set page margins (standard 1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        
    # --- PAGE 1: TITLE PAGE ---
    # Add spacing at the top
    for _ in range(3):
        doc.add_paragraph()
        
    # Acronym Title Block
    acronyms = [
        ("A", "utonomous", "자율형 에이전트"),
        ("M", "obile & Multimodal", "모바일 및 멀티모달"),
        ("E", "dge-centric", "엣지 중심 추론"),
        ("V", "oice & Vision", "음성 및 시각 물리 인지"),
        ("A", "rchitecture", "독립 실행형 아키텍처")
    ]
    
    for char, rest, desc in acronyms:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        
        # Big bold letter
        run_char = p.add_run(char)
        run_char.font.size = Pt(32)
        run_char.font.bold = True
        run_char.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Deep Blue
        
        # Acronym text
        run_rest = p.add_run(rest)
        run_rest.font.size = Pt(24)
        run_rest.font.bold = True
        run_rest.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        # Translation/Description
        run_desc = p.add_run(f"  ({desc})")
        run_desc.font.size = Pt(12)
        run_desc.font.italic = True
        run_desc.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        
    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(24)
    p_div.paragraph_format.space_after = Pt(24)
    run_div = p_div.add_run("____________________________________________________")
    run_div.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    
    # Portfolio Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = p_sub.add_run("AMEVA Projects Portfolio Summary\n")
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    run_dev = p_sub.add_run("Developer: ATSAdmin\nDate: June 2026")
    run_dev.font.size = Pt(10)
    run_dev.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_page_break()
    
    # --- PAGE 2: TABLE OF CONTENTS ---
    p_toc_title = doc.add_paragraph()
    p_toc_title.paragraph_format.space_before = Pt(24)
    p_toc_title.paragraph_format.space_after = Pt(24)
    run_toc_title = p_toc_title.add_run("Contents (목차)")
    run_toc_title.font.size = Pt(20)
    run_toc_title.font.bold = True
    run_toc_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    projects_list = [
        "1. 완전 자율형 로컬 멀티 에이전트 시스템 (AMEVA Agent Orchestra) (멀티 에이전트)",
        "2. 잠재적 인격 동역학 시뮬레이션 시스템 (AMEVA Dead Internet Theatre) (멀티 에이전트)",
        "3. 로컬 멀티모달 데스크탑 어시스턴트 (AMEVA Window Assistant) (혼합 에이전트)",
        "4. Whisper-Vosk 하이브리드 화자 분리 및 전사 엔진 (AMEVA STT Agent) (음성처리/STT)",
        "5. Windows CPU 최적화 Whisper LoRA 파인튜닝 플랫폼 (AMEVA STT Trainer) (음성처리/STT)",
        "6. 엣지-호스트 연동 및 물리 소거 동기화 파이프라인 (AMEVA Edge Agent) (음성처리/STT)",
        "7. 엣지 디바이스 기반 주기적 데이터 수집 및 자동화 파이프라인 (AMEVA Data Harvester) (음성처리/STT)",
        "8. 오프라인 및 로컬 지향 문서 처리 AI RAG 워크스테이션 (AMEVA Doc AI) (LLM)",
        "9. 컨테이너 격리식 고성능 LLM 벤치마킹 플랫폼 (AMEVA Benchmark Suite) (LLM)",
        "10. 분산형 LLM API 게이트웨이 및 워커 클러스터 (AMEVA Model Nexus) (LLM)"
    ]
    
    for proj in projects_list:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(8)
        run_item = p_item.add_run(proj)
        run_item.font.size = Pt(11)
        run_item.font.bold = True
        run_item.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    doc.add_page_break()
    
    # --- PAGE 3: PROJECT OVERVIEW ---
    p_ov_title = doc.add_paragraph()
    p_ov_title.paragraph_format.space_before = Pt(24)
    p_ov_title.paragraph_format.space_after = Pt(24)
    run_ov_title = p_ov_title.add_run("Project Overview (프로젝트 개요)")
    run_ov_title.font.size = Pt(20)
    run_ov_title.font.bold = True
    run_ov_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    overview_sections = [
        ("AMEVA 프로젝트 개요", 
         "Project AMEVA는 외부 클라우드 API 의존성 0%를 지향하며, 로컬/엣지 장비 단독 환경에서 작동 가능한 고성능 AI 에이전트 아키텍처를 실증하고 구현하는 프로젝트입니다. 'Beyond the Cloud, Intelligence Survives on the Edge'라는 슬로건 아래, 자본이나 대규모 서버가 아닌 최적화 및 엔지니어링 설계를 통해 독자적인 인공지능 자율 생태계를 구축하고자 합니다."),
        
        ("개발 철학 (Philosophy)", 
         "빅테크 기업의 클라우드 플랫폼에 의해 지능이 예속되는 시대를 거부하고, 엔지니어의 최적화 집념으로 독립 실행 가능한 '실행형 지능'의 민주화를 실현합니다. 네트워크 단절 및 극단적인 보안 통제 하에서도 기기 단독으로 온전히 동작하는 지적 생명체의 생존을 모사합니다."),
        
        ("보안성 (Security)", 
         "외부 네트워크 통신망을 물리적으로 차단하여 원천적인 중요 기밀 유출을 방지합니다. 또한, 전송 검증이 완료된 즉시 엣지 기기 내부의 원본 오디오 및 DB의 임시 흔적 데이터에 난수 비트를 덮어쓰는 포렌식 완전 소거(Data Shredding) 메커니즘을 상시 가동하여 장치 유출 위험을 차단합니다."),
        
        ("자율성 및 유기성 (Autonomy & Organicality)", 
         "각 컴포넌트(에이전트)는 기획(Architect), 구현(Developer), 검증(Tester) 등의 업무를 유기적으로 협업(Handoff)하여 처리합니다. 리소스 과부하 상황을 스스로 모니터링하여 OOM을 방지하는 SRE Watchdog과 결합되어, 환경 변화에 적응하며 지속 생존할 수 있는 통합 자율 지능 체계를 완성합니다."),

        ("개발 기여도 (Development Contribution)", 
         "본 AMEVA 프로젝트 포트폴리오에 수록된 10개 서브 시스템의 아키텍처 설계, 보안 격리 샌드박스, 이중 백업 파이프라인 및 가속 추론 코어 알고리즘 구현을 포함한 모든 기획·개발·검증 단계를 개발자 1인이 100% 단독으로 수행하였습니다.")
    ]
    
    for title, text in overview_sections:
        p_sec_title = doc.add_paragraph()
        p_sec_title.paragraph_format.space_before = Pt(12)
        p_sec_title.paragraph_format.space_after = Pt(4)
        run_sec_title = p_sec_title.add_run(title)
        run_sec_title.font.size = Pt(12)
        run_sec_title.font.bold = True
        run_sec_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        p_sec_text = doc.add_paragraph()
        p_sec_text.paragraph_format.space_after = Pt(12)
        p_sec_text.paragraph_format.line_spacing = 1.15
        run_sec_text = p_sec_text.add_run(text)
        run_sec_text.font.size = Pt(10)
        
    # Add Core Technology Stack
    p_stack_title = doc.add_paragraph()
    p_stack_title.paragraph_format.space_before = Pt(12)
    p_stack_title.paragraph_format.space_after = Pt(4)
    run_stack_title = p_stack_title.add_run("주요 기술 스택 (Main Technology Stack)")
    run_stack_title.font.size = Pt(12)
    run_stack_title.font.bold = True
    run_stack_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    tech_stacks = [
        ("Python", "백엔드 코어 및 비즈니스 로직 구현"),
        ("Docker", "에이전트 컨테이너 격리 및 독립 런타임 환경 구성"),
        ("llama.cpp", "소형 언어 모델 로컬 GPU/CPU 가속 추론"),
        ("Whisper", "고정밀 온디바이스 음성 인식 및 데이터 전사"),
        ("Vanilla JS", "실시간 관제 인터페이스 시각화 및 양방향 소켓 제어"),
        ("SQLite3", "에이전트 내부 상태 로그 및 대화 이력 로컬 영속화")
    ]
    
    for tech, desc in tech_stacks:
        p_tech_item = doc.add_paragraph()
        p_tech_item.paragraph_format.left_indent = Pt(12)
        p_tech_item.paragraph_format.space_after = Pt(4)
        r_tech = p_tech_item.add_run(f"- {tech}: ")
        r_tech.font.bold = True
        r_tech.font.size = Pt(10)
        r_desc = p_tech_item.add_run(desc)
        r_desc.font.size = Pt(10)
        
    doc.add_page_break()
    
    # --- PAGES 4+: PROJECT PROFILES (Re-ordered) ---
    projects_data = [
        {
            "num": 1,
            "title": "단일 8B 대형 모델 기반 롤플레잉 멀티 에이전트 협업 시스템 (AMEVA Agent Orchestra) (멀티 에이전트)",
            "purpose": "기존 대형 언어 모델 에이전트의 두 가지 근본적 문제(클라우드 보안 리스크, 멀티 에이전트 구성 시 컨텍스트 오염 및 제어 불가능성)를 해결하고자, 단일 8B 모델을 상주시켜 메모리를 보존하면서 턴 단위 역할 주입과 데이터베이스 기반 메모리 격리를 통해 멀티 에이전트 협업을 시뮬레이션하는 자율 오케스트레이션 시스템입니다.",
            "features": "기존 에이전트의 내부 의사결정 및 프롬프트 구조를 리버스 엔지니어링 분석하여 도출한 구조를 토대로, 프로젝트 매니저의 승인 검토 루프를 통한 설계자, 개발자, 테스터 간의 요약 전달 협업을 지원합니다. 또한 보안 가이드라인 위반 시 작업을 중단하지 않고 피드백을 전달하여 안전한 코드를 자율 재생성하는 우회 복구 흐름을 탑재했습니다.",
            "tech": "Python, FastAPI, uvicorn, jinja2, psutil, GPUtil, sqlite3, llama-cpp-python, watchdog",
            "models": "Llama-3.1-8B-Instruct-GGUF 등 단일 8B 대형 모델",
            "algo": "LlamaGrammar 기능을 활용해 추론 단계에서 JSON 문법을 필터링하는 스키마 제어 알고리즘, 불완전 JSON의 후행 괄호를 탐색하여 자동 회생시키는 StrictParser, SQLite3 기반의 에이전트별 독립 장기 기억 관리 체계.",
            "control": "단일 인스턴스의 턴 기반 호출 흐름을 통제하는 뮤텍스 락 제어, 중괄호 스택 매칭 기반의 2차 JSON 복구 필터, 작업 공간 이탈 방지 경로 검증 및 eval, exec, subprocess, os.remove 등의 위험 명령어를 실시간 스캔 및 차단하는 샌드박스 보안 가드 모듈.",
            "achieve": "역할 격리를 통해 정체성 이탈 및 지시어 유출 문제를 방어하고, 에러 피드백 루프를 결합하여 자율 디버깅 성공률 92.0%를 정량 달성했습니다. 또한 LlamaGrammar 스키마 가드를 연동해 구문 규격 붕괴율 0% 및 10턴 연속 페르소나 유지를 검증하여 온프레미스 자율 생존 지능 체계의 안정성을 실증하였습니다."
        },
        {
            "num": 2,
            "title": "잠재적 인격 동역학 시뮬레이션 시스템 (AMEVA Dead Internet Theatre) (멀티 에이전트)",
            "purpose": "페르소나와 감정 벡터가 주입된 다수의 자율 AI 에이전트들이 폐쇄된 가상 포럼 공간에서 상호 작용할 때 발생하는 전체 여론 동역학의 수렴 및 전이 양상을 규명하기 위한 연구 프로젝트입니다.",
            "features": "에이전트별 감정, 의견, 영향력 지표의 LPDE 다차원 상태 벡터 업데이트, 비결정적 대화 행동 (Ignore, Reply, Join) 선택 알고리즘, 대화 정체 해소를 위한 감독 에이전트 외란 주입 및 실시간 모니터링 시뮬레이터 가동을 포함합니다.",
            "tech": "Python, FastAPI, Docker Engine SDK, SQLite3, TailWindCSS, Chart.js",
            "models": "Qwen2.5 (0.5B/3B), Llama-3.2 (1B), Llama-3.1 (8B)",
            "algo": "감정 (Affect) 2D, 의견 (Opinion) 4D, 영향력 (Power) 2D 텐서곱으로 상태를 정의하는 LPDE 엔진 수식, 대화 엣지 관계 업데이트용 지수이동평균 (EMA) 필터링, 하드라이너 자아 붕괴를 정규식으로 감지해 무효화하는 Stance Coherence 검증 로직, 싱글 멘션 강제 정제 알고리즘.",
            "control": "거대 LLM의 화두 발의 후 독립된 Docker 컨테이너 격리 런타임 내 봇 에이전트 기동, 감독 에이전트의 델타 개입 및 예외 상황 감지 시 중 Neutral 텐서 Fallback 제어 흐름.",
            "achieve": "감독 에이전트 부재 시 여론이 단일 극단으로 수렴되는 한계를 분석하였으며, 소형 모델의 정체성 붕괴 문제를 발견하고 분노/공감 수치 벡터화를 통해 인간과 유사한 창발적 상호작용 모형을 실증하였습니다."
        },
        {
            "num": 3,
            "title": "로컬 멀티모달 데스크탑 어시스턴트 (AMEVA Window Assistant) (혼합 에이전트)",
            "purpose": "Copilot 및 Gemini 화면공유 AI 비서 시스템을 로컬/온프레미스 단독 환경에서 가동하기 위해 설계되었습니다. 외부 망 연결 없이 데스크탑 화면 및 음성 입력을 받아 시각적 의도를 분석하고 오디오 합성 피드백을 제공함으로써 하드웨어 한계 내 실시간 보조 수행을 검증합니다.",
            "features": "다중 모니터 DPI 인식 화면 프레임 덤프, 16kHz PCM 실시간 오디오 캡처, Tesseract OCR을 통한 텍스트 좌표계 Scene Graph화, 의도 라우팅(Intent Routing) 기반 분기 제어 및 Windows SAPI 음성 합성 출력을 제공합니다.",
            "tech": "Python, Tesseract OCR, sounddevice, wave, mss, Windows SAPI",
            "models": "Whisper.cpp (Small/Tiny), Qwen2.5 (1.5B/Router), Qwen2-VL (2B/VLM), Llama-3.1 (8B/LLM)",
            "algo": "화면의 텍스트 밀도와 질문 패턴을 파악하여 OCR-LLM 경로와 VLM 이미지 경로로 실시간 스위칭하는 하이브리드 의도 라우팅 알고리즘, 사용자 단축키 및 음성 명령 트리거를 통해 수집된 윈도우 PCM 스트림의 RMS 에너지를 감시하여 가청 임계치 미달 시 전송을 자동 차단하는 무음 감지 로직, TTS 음성 낭독을 위해 Chain-of-Thought 메타데이터와 마크다운 기호를 Stripping하는 텍스트 클리닝 필터.",
            "control": "단축키 및 음성 명령 입력으로부터 OS PCM 스트림 수집 및 RMS 감시 차단, 전사 텍스트의 intent_router 분기 판정 및 VLM/OCR 매핑 추론, 결과 텍스트 SAPI 전달 및 SQLite3 이력 로깅 처리 흐름.",
            "achieve": "로컬 단독 가동 실증 과정에서 컴포넌트 동시 구동 시 12~14GB VRAM 점유 및 약 5~8초의 End-to-End 지연이 유발됨을 수치 검증하여 저성능 컴퓨터 환경 하의 물리적 가용 한계를 정밀 규명하였습니다."
        },
        {
            "num": 4,
            "title": "Whisper-Vosk 하이브리드 화자 분리 및 전사 엔진 (AMEVA STT Agent) (음성처리/STT)",
            "purpose": "보안이 극대화된 폐쇄 환경에서 외부 API 전송 없이 고정밀 음성 인식 및 다중 화자 분류(Diarization)를 실현하고, 분석 데이터를 태스크 배치 단위로 자동 캡슐화 및 시각 보고서로 발행하는 대시보드 플랫폼입니다.",
            "features": "Whisper 전사 및 Vosk 화자 임베딩 결합, K-Means/PCA 기반 화자 군집 시각화 차트 제공, python-docx 활용 엔터프라이즈 회의록 워드 보고서 즉시 출력, SQLite 분석 이력 데이터 백로그 관리를 포함합니다.",
            "tech": "Python, Streamlit, pywhispercpp, vosk, scikit-learn, Plotly, pandas, python-docx",
            "models": "Whisper.cpp (Small/Tiny/GGUF), Vosk STT Model, Vosk Spk Model",
            "algo": "시간축 오차가 있는 두 다른 음향 특징을 정합하기 위해 세그먼트 중앙값 기준 최단 거리에 있는 화자 클러스터를 1:1 강제 매핑하는 Shortest Distance Mapping (SDM) 알고리즘, X-Vector 코사인 유사도 산출 및 PCA 축소 K-Means 군집 분석.",
            "control": "GIL에 의한 화면 동결 방지를 위해 무거운 음향 처리를 multiprocessing.Process 하위 OS 프로세스로 격리하고, 비차단 비동기 큐를 통해 대시보드 화면에 로그 스트림을 브로드캐스팅하는 흐름.",
            "achieve": "sys.platform bypass를 통해 모바일 터미널 환경 동작을 검증하고, SDM 매핑 기법 도입으로 단락 화자 지정 정확도를 비약적으로 상향하였습니다. 또한, Dead Internet Theatre 시뮬레이션 합성 대화 오디오 검증을 통해 WER 3.2%의 높은 전사 신뢰도 및 4개 에이전트 인격 중심점의 2D 평면상 완벽 분리를 정량 실증하였습니다."
        },
        {
            "num": 5,
            "title": "Windows CPU 최적화 Whisper LoRA 파인튜닝 플랫폼 (AMEVA STT Trainer) (음성처리/STT)",
            "purpose": "척박한 윈도우 CPU 환경에서 시사/경제 등 특정 도메인에 특화된 Whisper 모델 학습을 위해 자막 싱크 중복을 제거하고 LoRA 파인튜닝을 기동하며, 훈련 후 가중치를 원본 모델과 무손실 병합하여 배포용 GGUF 파일로 변환 가속하는 플랫폼입니다.",
            "features": "접두-접미사 중복 제거 파싱, 자막 종결어미 문맥 보호 dynamic chunking, IterableDataset 기반 훈련 스트리밍, LoRA 어댑터 가중치 병합 및 4-bit 양자화 컴파일 연계를 제공합니다.",
            "tech": "Python, PyTorch (FP32 CPU), transformers, PEFT (LoRA), datasets, librosa, soundfile, pydub, python-docx, WandB",
            "models": "Whisper (Tiny/Small/Medium)",
            "algo": "대규모 오디오 로드 시의 WinError 87 한계를 우회하기 위한 IterableDataset 실시간 스트리밍 적재 알고리즘, 공백 제거 단어 기반 접미사-접두사 중복 매칭 제거(Suffix-Prefix Overlap Matching), 학습 전 구어체 연속 1~3 gram 반복 및 중복 데이터 validator 품질 감사 알고리즘.",
            "control": "가상환경 직렬화 Pickle Error 차단을 위해 제너레이터와 로거 객체 생명주기를 메모리 레벨에서 독립 분리하며, uvicorn API 서버(8600 포트) 상태를 검사해 꺼져 있을 시 백그라운드 프로세스로 자동 serve 구동합니다.",
            "achieve": "CPU 단독 윈도우 환경 내에서 OOM 실패 없이 LoRA 훈련 무중단 신뢰성을 확보하고, whisper.cpp 양자화 연계로 이식 속도를 극대화했습니다."
        },
        {
            "num": 6,
            "title": "엣지-호스트 연동 및 물리 소거 동기화 파이프라인 (AMEVA Edge Agent) (음성처리/STT)",
            "purpose": "기밀성이 요구되는 모바일 엣지(Termux) 환경에서 오디오 녹음 및 STT, 요약을 수행하고, 이관 스케줄에 맞춰 호스트 워크스테이션으로 백업한 직후 로컬 기기의 모든 원본 및 임시 흔적 데이터를 영구 삭제하는 보안 전송 파이프라인입니다.",
            "features": "파이썬 표준 라이브러리 100% 가동 엣지 스캐너, 호스트 감시(Watchdog) 데몬의 실시간 DB 병합 및 무결성 검증, 데이터 전송 성공 확인 즉시 로컬 데이터 완전 삭제 프로세스를 가동합니다.",
            "tech": "Python 3.x (엣지 표준 모듈 100%), watchdog, reportlab, requests, sqlite3, subprocess",
            "models": "Whisper.cpp (Small/STT), Llama-3 (8B/BitNet 1.58b), Llama-3.2 (3B/LLM)",
            "algo": "디렉토리 파일 I/O 감시를 위한 watchdog 기반 인입 이벤트 핸들러, 파일 복구를 물리적으로 저지하기 위해 무작위 난수 비트로 대상 섹터를 덮어쓰고 캐시 동기화를 집행하는 shred_file 및 shred_database 완전 소거 알고리즘, 1:1:1 물리 파일 크기/해시 교차 검증 로직.",
            "control": "엣지 오디오 스캔 및 데이터 DB 기록 -> STT/요약 대기 -> SCP 파일 전송 및 크기 검증 -> shred_file 기동 -> 마이그레이션용 복제 DB 전송 -> 호스트 watchdog 감지 후 Master DB 병합 트랜잭션 수행 -> 성공 시그널 엣지 수신 및 원본 DB 완전 제거의 제어 흐름.",
            "achieve": "엣지 프로그램의 외부 종속성을 완전히 제거하여 모바일 이식성을 높였으며, 파일 포렌식 복구를 무력화시키는 완전 삭제 기법을 실증하였습니다."
        },
        {
            "num": 7,
            "title": "엣지 디바이스 기반 주기적 데이터 수집 및 자동화 파이프라인 (AMEVA Data Harvester) (음성처리/STT)",
            "purpose": "안드로이드 Termux 및 Windows 환경의 디렉토리를 상시 감시하여 신규 인입되는 음성 녹음 데이터를 자동 감지하고, 로컬 Whisper STT 및 소형 SLM을 통해 요약 및 번역 가공을 완수하여 메인 서버로 안전하게 자동 전송하는 파이프라인입니다.",
            "features": "시간차 폴링 기반 파일 감지, 로컬 STT 전사 및 요약 가공, PAC 다중 통신망 기반 동적 전송 라우팅, ZIP 해시 검증 및 전송 완료 원본 파일 안전 소거를 지원합니다.",
            "tech": "Python, standard libraries, requests, zipfile",
            "models": "Whisper (Small/STT), Qwen (1.8B), Phi-3 (3B), Llama-3.1 (8B)",
            "algo": "파일 쓰기 락 상태를 검증하기 위해 이중 시간차 스냅샷을 생성하여 크기와 수정 시간이 정체된 시점을 안정화 상태로 파악하는 O(1) 배치 폴링 검증 알고리즘, Primary(SSH/SCP) -> Alternate(API) -> Contingency(Telegram)로 전환하는 PAC 라우팅 알고리즘, SHA-256 해시 검증.",
            "control": "파일 감지 -> 로컬 전사/요약 가공 -> ZIP 압축 및 SHA-256 해시 생성 -> PAC 복합 전송 트리거 -> 전송 성공 확인 후 원본 데이터 로컬 소거 프로세스 흐름.",
            "achieve": "엣지 디바이스 환경에서 오프라인 요약 전처리 후 가공 텍스트 위주 전송으로 서버 대역폭 부하를 100분의 1 수준으로 최소화하고 기기 사양별 한계 측정을 실증하였습니다."
        },
        {
            "num": 8,
            "title": "오프라인 및 로컬 지향 문서 처리 AI RAG 워크스테이션 (AMEVA Doc AI) (LLM)",
            "purpose": "중요 정보 보안이 핵심인 환경에서 외부 네트워크 연결을 차단하고, 로컬 디바이스 단독으로 오프라인 문서 파싱, 청크별 병렬 분산 요약, RAG 기반 질의응답 및 TTS 낭독 오디오북 제작을 서비스하는 워크스테이션입니다.",
            "features": "HWP, DOCX, XLSX 등 실무 문서 텍스트 마크다운 정합 추출, 스레드 가중치 기반 분산 병렬 요약 처리, RAG 기반 지식 검색 QA, Edge-TTS 오디오북 낭독 에셋 생성을 제공합니다.",
            "tech": "Python, PyQt6, ReportLab, OleFile, python-docx, openpyxl, edge-tts, websockets",
            "models": "Qwen2.5 (1.5B/3B), Gemma2 (2B), Llama-3.1 (8B)",
            "algo": "문서 확장자별 정규식 전처리 및 가변 청킹(Chunking) 알고리즘, ThreadPoolExecutor를 사용한 비동기 분산 요약 병합 로직, 메모리 경합 방지용 리소스 락킹 제어.",
            "control": "기기 전력 상황을 감시하는 경찰 워커(PoliceWorker)를 기동하여 노트북 배터리 방전 상황 감지 즉시 멀티스레드 작업을 강제 축소하고, 문서 요약과 RAG 채팅 스레드가 충돌하지 않도록 트랜잭션을 락킹 제어합니다.",
            "achieve": "오프라인 로컬 대규모 문서 요약을 완수해 정보 로컬화 보안을 확보하였으며, 스레드 가용 개수에 따른 속도와 메모리 안정성 간의 트레이드오프 관계를 분석해 최적 조절 비율을 수립하였습니다."
        },
        {
            "num": 9,
            "title": "컨테이너 격리식 고성능 LLM 벤치마킹 플랫폼 (AMEVA Benchmark Suite) (LLM)",
            "purpose": "엣지 디바이스 환경에서 파편화된 자원 조건에 대응하여, 모델별 실질 추론 성능과 전력 효율(Tokens/J), 지식 정합성을 독립 격리 컨테이너 환경에서 측정하고 정량 보고서를 도출하는 플랫폼입니다.",
            "features": "Docker Compose 가상 아레나를 통한 모델 런타임 물리 격리, 하네스 태스크 정규식 자동 채점, 로컬 Ollama 연동 8B 이상 AI 판정관(Judge) 정성 평가, 진단 워드(.docx)/엑셀 보고서 자동 생성을 담당합니다.",
            "tech": "Python, FastAPI, Docker SDK, SQLite3, python-docx, openpyxl, Plotly",
            "models": "EXAONE-3.5 (7.8B/AI 판정관), Llama-3.1 (8B), Qwen2.5 (1.5B/3B/7B)",
            "algo": "예상 정규 표현식을 대조하여 정밀한 정량 판정을 내리는 자동 채점 알고리즘, 리소스 누수를 막기 위해 로그 히스토리를 500라인 원형 큐 버퍼로 중계하는 비동기 브로드캐스터, 판정 결과 Dirty JSON 정밀 정제기.",
            "control": "모델 추론 시 격리 Cgroup 하한을 통제하며, 벤치마크 모델과 판정관 모델이 메모리에 동시에 상주하지 않도록 순차 언로드 수명 주기를 강제합니다.",
            "achieve": "모델 전환 시 기존 컨테이너를 파괴하고 클린 스테이트 상태를 회복하는 Smart SWAP 기법을 구현하여 계측 지표의 100% 결정성 및 재현 가능성을 확보하였으며, VRAM 경합 크래시를 완벽 차단하였습니다."
        },
        {
            "num": 10,
            "title": "분산형 LLM API 게이트웨이 및 워커 클러스터 (AMEVA Model Nexus) (LLM)",
            "purpose": "단일 노드의 물리적 하드웨어 한계를 극복하기 위해 다중 Docker 컨테이너 워커(Worker)와 라우팅 게이트웨이(Router)를 연동한 무중단 분산 LLM 서빙 플랫폼입니다. 큐(Queue) 버퍼링 메커니즘을 적용하여 모델 교체(Hot-swap)나 노드 장애 시에도 요청 연결 유실률 0%를 달성하는 고신뢰성 라우팅 설계를 구현합니다.",
            "features": "무중단 동적 가중치 스와핑(Hot-swap) 제어, 비동기 큐 버퍼링 및 장애 격리(PENDING 대기열), 하드웨어 상태(CUDA 유무) 탐지 기반 지능형 분산 로드밸런서 제공을 목표로 합니다.",
            "tech": "Python, FastAPI, Docker Compose, SQLite (WAL Mode), Asyncio Queueing",
            "models": "Llama-3-8B-GGUF, Qwen2.5-3B-GGUF",
            "algo": "비동기 큐 기반의 무중단 핫스왑(Hot-Swap) 및 동적 스왑 지연 방어 메커니즘, Python gc.collect() 강제를 통한 VRAM 누수 및 단편화 방지 알고리즘, Hardware-Aware 동적 디스패칭 로드밸런서.",
            "control": "Gateway 인입 -> SQLite Queue 버퍼 상태 전이 -> 대상 워커 Poll 작업 할당 -> SSE(Server-Sent Events) 스트리밍 토큰 반환 -> 완료 시 WAL 로거 아티팩트 보존 흐름 제어.",
            "achieve": "도커 격리 기반의 무중단 모델 교체 파이프라인을 구축하여 VRAM 오버플로우와 모델 교체 시 커넥션 끊김 문제를 하드웨어 단독 서버 제약 하에서 완벽히 해결함을 실증하였습니다."
        }
    ]
    
    for proj in projects_data:
        # Heading 1: Project Title
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(24)
        p_title.paragraph_format.space_after = Pt(12)
        run_title = p_title.add_run(f"{proj['num']}. {proj['title']}")
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        # 1. 주제 및 목적
        p_h2_1 = doc.add_paragraph()
        p_h2_1.paragraph_format.space_before = Pt(12)
        p_h2_1.paragraph_format.space_after = Pt(4)
        run_h2_1 = p_h2_1.add_run("주제 및 목적")
        run_h2_1.font.size = Pt(12)
        run_h2_1.font.bold = True
        run_h2_1.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        p_txt_1 = doc.add_paragraph()
        p_txt_1.paragraph_format.space_after = Pt(12)
        p_txt_1.paragraph_format.line_spacing = 1.15
        run_txt_1 = p_txt_1.add_run(proj["purpose"])
        run_txt_1.font.size = Pt(10)
        
        # 2. 주요 기능 및 연구 목표
        p_h2_2 = doc.add_paragraph()
        p_h2_2.paragraph_format.space_before = Pt(12)
        p_h2_2.paragraph_format.space_after = Pt(4)
        run_h2_2 = p_h2_2.add_run("주요 기능 및 연구 목표")
        run_h2_2.font.size = Pt(12)
        run_h2_2.font.bold = True
        run_h2_2.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        p_txt_2 = doc.add_paragraph()
        p_txt_2.paragraph_format.space_after = Pt(12)
        p_txt_2.paragraph_format.line_spacing = 1.15
        run_txt_2 = p_txt_2.add_run(proj["features"])
        run_txt_2.font.size = Pt(10)
        
        # 3. 기술 스택 및 핵심 알고리즘
        p_h2_3 = doc.add_paragraph()
        p_h2_3.paragraph_format.space_before = Pt(12)
        p_h2_3.paragraph_format.space_after = Pt(4)
        run_h2_3 = p_h2_3.add_run("기술 스택 및 핵심 알고리즘")
        run_h2_3.font.size = Pt(12)
        run_h2_3.font.bold = True
        run_h2_3.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        p_tech = doc.add_paragraph()
        p_tech.paragraph_format.space_after = Pt(4)
        p_tech.paragraph_format.left_indent = Pt(12)
        r_tech_lbl = p_tech.add_run("사용 기술: ")
        r_tech_lbl.font.bold = True
        p_tech.add_run(proj["tech"])
        
        p_models = doc.add_paragraph()
        p_models.paragraph_format.space_after = Pt(8)
        p_models.paragraph_format.left_indent = Pt(12)
        r_models_lbl = p_models.add_run("사용 모델: ")
        r_models_lbl.font.bold = True
        p_models.add_run(proj["models"])
        
        p_algo = doc.add_paragraph()
        p_algo.paragraph_format.space_after = Pt(8)
        p_algo.paragraph_format.left_indent = Pt(12)
        p_algo.paragraph_format.line_spacing = 1.15
        r_algo_lbl = p_algo.add_run("핵심 알고리즘: ")
        r_algo_lbl.font.bold = True
        p_algo.add_run(proj["algo"])
        
        p_control = doc.add_paragraph()
        p_control.paragraph_format.space_after = Pt(12)
        p_control.paragraph_format.left_indent = Pt(12)
        p_control.paragraph_format.line_spacing = 1.15
        r_control_lbl = p_control.add_run("에이전트 및 보안 제어: ")
        r_control_lbl.font.bold = True
        p_control.add_run(proj["control"])
        
        for p_el in [p_tech, p_models, p_algo, p_control]:
            for run in p_el.runs:
                run.font.size = Pt(10)
        
        # 4. 연구 성과
        p_h2_4 = doc.add_paragraph()
        p_h2_4.paragraph_format.space_before = Pt(12)
        p_h2_4.paragraph_format.space_after = Pt(4)
        run_h2_4 = p_h2_4.add_run("연구 성과")
        run_h2_4.font.size = Pt(12)
        run_h2_4.font.bold = True
        run_h2_4.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        p_ach = doc.add_paragraph()
        p_ach.paragraph_format.space_after = Pt(12)
        p_ach.paragraph_format.line_spacing = 1.15
        run_ach_lbl = p_ach.add_run("연구 성과: ")
        run_ach_lbl.font.bold = True
        p_ach.add_run(proj["achieve"])
        
        for p_el in [p_ach]:
            for run in p_el.runs:
                run.font.size = Pt(10)
                
        # Page break after each project except the last
        if proj["num"] < 10:
            doc.add_page_break()
            
    # Save document
    doc.save(doc_path)
    print("Portfolio docx generated successfully at:", doc_path)

if __name__ == '__main__':
    create_portfolio()
