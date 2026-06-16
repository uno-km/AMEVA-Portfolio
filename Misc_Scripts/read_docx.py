import docx
import os
import sys

doc_path = r"c:\ameva\ameva\combined_profile_summary.docx"
if not os.path.exists(doc_path):
    print("File not found:", doc_path)
    exit(1)

doc = docx.Document(doc_path)
output_path = r"c:\ameva\docx_content.txt"

with open(output_path, "w", encoding="utf-8") as f:
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
    f.write(f"Total tables: {len(doc.tables)}\n\n")
    f.write("--- ALL PARAGRAPHS ---\n")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            f.write(f"P{i} [{p.style.name}]: {p.text}\n")
            
    f.write("\n--- TABLES OUTLINE ---\n")
    for i, table in enumerate(doc.tables):
        f.write(f"Table {i}: rows={len(table.rows)}, cols={len(table.columns)}\n")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                f.write(f"  Cell ({r_idx}, {c_idx}): {cell.text}\n")

print("Done writing to docx_content.txt")
