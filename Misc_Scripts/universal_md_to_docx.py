import argparse
import os
import re
import base64
import requests
import docx
from io import BytesIO
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '  <w:left w:val="none"/>'
        '  <w:right w:val="none"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '  <w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def download_mermaid_image(mermaid_code):
    encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
    url = f'https://mermaid.ink/img/{encoded}'
    try:
        response = requests.get(url, timeout=15)
        if response.status_code == 200:
            return BytesIO(response.content)
    except Exception as e:
        print(f"Error downloading mermaid chart: {e}")
    return None

def parse_inline_markdown(paragraph, text, is_quote=False):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            subparts = re.split(r'(\*.*?\*)', part)
            for sub in subparts:
                if sub.startswith('*') and sub.endswith('*'):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    run = paragraph.add_run(sub)
                
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(10)
                if is_quote:
                    run.font.color.rgb = RGBColor(100, 100, 100)
                    run.italic = True
        
        for r in paragraph.runs:
            r.font.name = 'Malgun Gothic'

def process_table(doc, table_lines):
    cleaned_lines = []
    for line in table_lines:
        if not line.strip() or re.match(r'^[\s\|\-\:]+$', line):
            continue
        cols = [c.strip() for c in line.strip().strip('|').split('|')]
        cleaned_lines.append(cols)
        
    if not cleaned_lines:
        return

    table = doc.add_table(rows=len(cleaned_lines), cols=max(len(r) for r in cleaned_lines))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_borders(table)

    for r_idx, row_data in enumerate(cleaned_lines):
        row = table.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            if c_idx >= len(row.cells):
                break
            cell = row.cells[c_idx]
            
            clean_text = text.replace('**', '') 
            cell.text = clean_text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            p = cell.paragraphs[0]
            
            if r_idx == 0:
                set_cell_background(cell, "003366")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Malgun Gothic'
                    r.font.size = Pt(9.5)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F7F9FB")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Malgun Gothic'
                    r.font.size = Pt(9)
                    if text.startswith('**') and text.endswith('**'):
                        r.font.bold = True

def convert_md_to_docx(input_md, output_docx):
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)

    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_mermaid = False
    mermaid_code = []
    in_table = False
    table_lines = []

    is_first_h1 = True
    is_subtitle = False

    for line in lines:
        raw_line = line.strip('\n')
        stripped = raw_line.strip()

        if stripped.startswith('```mermaid'):
            in_mermaid = True
            continue
        elif in_mermaid and stripped.startswith('```'):
            in_mermaid = False
            img_io = download_mermaid_image('\n'.join(mermaid_code))
            if img_io:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_io, width=Inches(5.5))
            mermaid_code = []
            continue
        elif in_mermaid:
            mermaid_code.append(raw_line)
            continue
            
        if stripped.startswith('```'):
            continue

        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            process_table(doc, table_lines)
            in_table = False
            table_lines = []
            if not stripped:
                continue

        if not stripped:
            continue

        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            img_path = img_match.group(2)
            input_dir = os.path.dirname(os.path.abspath(input_md))
            portfolio_dir = r'c:\ameva\AMEVA-Portfolio'
            
            if os.path.isabs(img_path):
                full_img_path = img_path
            else:
                full_img_path = os.path.join(input_dir, img_path)
                if not os.path.exists(full_img_path):
                    full_img_path = os.path.join(portfolio_dir, img_path)
            
            if os.path.exists(full_img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                p.add_run().add_picture(full_img_path, width=Inches(5.5))
                
                caption_p = doc.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.paragraph_format.space_after = Pt(8)
                run = caption_p.add_run(f"▲ {img_match.group(1)}")
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(102, 102, 102)
                run.italic = True
            continue

        if stripped.startswith('#'):
            level = len(stripped.split()[0])
            text = stripped[level:].strip()
            
            if level == 1 and is_first_h1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                is_first_h1 = False
                is_subtitle = True
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                is_subtitle = False
            continue

        if is_subtitle and not stripped.startswith('>'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clean_sub = stripped.strip('*')
            run = p.add_run(clean_sub)
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(100, 100, 100)
            is_subtitle = False
            continue

        is_subtitle = False

        if stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parse_inline_markdown(p, stripped[1:].strip(), is_quote=True)
            continue
            
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_markdown(p, stripped[2:].strip())
            continue

        p = doc.add_paragraph()
        
        num_match = re.match(r'^(\d+\)|①|②|③|④|⑤)\s*(.*)', stripped)
        if num_match:
            r_num = p.add_run(num_match.group(1) + " ")
            r_num.font.name = 'Malgun Gothic'
            r_num.font.size = Pt(12)
            r_num.font.bold = True
            parse_inline_markdown(p, num_match.group(2))
        else:
            parse_inline_markdown(p, stripped)

    if in_table and table_lines:
        process_table(doc, table_lines)

    doc.save(output_docx)
    print(f"Successfully converted {input_md} to {output_docx}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Universal MD-to-DOCX Converter (AMEVA Corporate Theme)")
    parser.add_argument("--input", required=True, help="Path to input Markdown file")
    parser.add_argument("--output", required=True, help="Path to output DOCX file")
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Error: Input file {args.input} does not exist.")
        exit(1)
        
    convert_md_to_docx(args.input, args.output)
