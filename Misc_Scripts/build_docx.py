# -*- coding: utf-8 -*-
import os
import re
import codecs
import base64
import glob
import requests
from PIL import Image

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

PORTFOLIO_DIR = r'c:\ameva\AMEVA-Portfolio'
IMG_DIR = os.path.join(PORTFOLIO_DIR, 'img')

if not os.path.exists(IMG_DIR):
    os.makedirs(IMG_DIR)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def parse_bold_runs(paragraph, text, is_italic=False, font_name='Malgun Gothic', font_size=10.5, font_color=RGBColor(51, 51, 51)):
    # Split text by bold markers (**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.italic = is_italic
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = font_color

def download_mermaid_image(mermaid_code, output_name):
    encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
    url = f'https://mermaid.ink/img/{encoded}'
    try:
        response = requests.get(url, timeout=15)
        if response.status_code == 200:
            png_path = os.path.join(IMG_DIR, f'{output_name}.png')
            jpg_path = os.path.join(IMG_DIR, f'{output_name}.jpg')
            with open(png_path, 'wb') as f:
                f.write(response.content)
            # Convert to JPG
            im = Image.open(png_path)
            im.convert('RGB').save(jpg_path, 'JPEG')
            # Clean up PNG
            try:
                os.remove(png_path)
            except:
                pass
            return f'img/{output_name}.jpg'
    except Exception as e:
        print(f"Error fetching mermaid image ({output_name}): {e}")
    return None

def preprocess_markdown(md_path, base_name):
    with codecs.open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Find and extract all mermaid blocks
    mermaid_pattern = re.compile(r'```mermaid\s*(.*?)\s*```', re.DOTALL)
    matches = list(mermaid_pattern.finditer(content))
    
    # Process from last to first to keep string offsets valid
    for idx, match in enumerate(reversed(matches)):
        mermaid_code = match.group(1)
        img_id = len(matches) - idx
        img_name = f'{base_name}_chart_{img_id}'
        
        # Fetch image
        img_relative_path = download_mermaid_image(mermaid_code, img_name)
        if img_relative_path:
            # Replace the mermaid block in the markdown content with image markdown link
            start, end = match.span()
            content = content[:start] + f'![시스템 흐름도]({img_relative_path})' + content[end:]
            
    # Remove any horizontal rules '---'
    lines = content.split('\n')
    cleaned_lines = []
    for line in lines:
        if line.strip() == '---':
            continue
        cleaned_lines.append(line)
        
    final_content = "\n".join(cleaned_lines)
    
    # Overwrite the markdown file to keep it clean and updated
    with codecs.open(md_path, 'w', encoding='utf-8') as f:
        f.write(final_content)
        
    return final_content

def render_table(doc, table_rows):
    if not table_rows:
        return
    
    # Parse table rows into 2D list
    matrix = []
    for row in table_rows:
        cols = [col.strip() for col in row.split('|')[1:-1]]
        matrix.append(cols)
        
    if not matrix:
        return
        
    num_rows = len(matrix)
    num_cols = len(matrix[0])
    
    # Skip separator row if present (e.g. |---|---|)
    if num_rows > 1 and all(re.match(r'^:?-+:?$', col) for col in matrix[1]):
        matrix.pop(1)
        num_rows -= 1

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Apply cell widths evenly
    for row_idx, row_data in enumerate(matrix):
        row = table.rows[row_idx]
        is_header = (row_idx == 0)
        bg_color = '1B365D' if is_header else ('F2F5F8' if row_idx % 2 == 1 else 'FFFFFF')
        
        for col_idx, text in enumerate(row_data):
            # Avoid out of index if row has fewer cells
            if col_idx >= len(row.cells):
                continue
            cell = row.cells[col_idx]
            cell.text = "" # Clear default
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # Formatting text
            text_color = RGBColor(255, 255, 255) if is_header else RGBColor(51, 51, 51)
            parse_bold_runs(p, text, font_size=9.5, font_color=text_color)
            if is_header:
                for run in p.runs:
                    run.bold = True
                    
            # Set Margins & Background
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_background(cell, bg_color)

def build_single_docx(md_path, docx_path, base_name):
    # Preprocess markdown (mermaid images, horizontal lines)
    markdown_text = preprocess_markdown(md_path, base_name)
    
    doc = Document()
    
    # Margins setup (54 Pt)
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)
        
    lines = markdown_text.split('\n')
    
    # State machine variables
    in_code_block = False
    code_lines = []
    
    in_table = False
    table_lines = []
    
    idx = 0
    num_lines = len(lines)
    
    while idx < num_lines:
        line = lines[idx]
        stripped = line.strip()
        
        # 1. Code block handling
        if stripped.startswith('```') and not stripped.startswith('```mermaid'):
            if in_code_block:
                # Close code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                
                code_text = "\n".join(code_lines)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(68, 68, 68)
                
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            idx += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue
            
        # 2. Table handling
        if stripped.startswith('|'):
            in_table = True
            table_lines.append(line)
            idx += 1
            continue
        else:
            if in_table:
                # Process table before moving on
                render_table(doc, table_lines)
                table_lines = []
                in_table = False
                
        # Skip empty lines
        if stripped == '':
            idx += 1
            continue
            
        # 3. Image tag handling: ![alt](path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            img_path = img_match.group(2)
            # Support both relative and absolute paths
            full_img_path = img_path if os.path.isabs(img_path) else os.path.join(PORTFOLIO_DIR, img_path)
            
            if os.path.exists(full_img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                p.add_run().add_picture(full_img_path, width=Inches(5.5))
                
                # Image caption
                caption_p = doc.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.paragraph_format.space_after = Pt(8)
                run = caption_p.add_run(f"▲ {img_match.group(1)}")
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(102, 102, 102)
                run.italic = True
            else:
                print(f"Warning: Image file not found at {full_img_path}")
            idx += 1
            continue
            
        # 4. Header 1 (# )
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            parse_bold_runs(p, stripped[2:], font_size=18, font_color=RGBColor(27, 54, 93))
            for r in p.runs:
                r.bold = True
            idx += 1
            continue
            
        # 5. Header 2 (## )
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            parse_bold_runs(p, stripped[3:], font_size=14, font_color=RGBColor(27, 54, 93))
            for r in p.runs:
                r.bold = True
            idx += 1
            continue
            
        # 6. Header 3 (### )
        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            parse_bold_runs(p, stripped[4:], font_size=12, font_color=RGBColor(46, 91, 136))
            for r in p.runs:
                r.bold = True
            idx += 1
            continue
            
        # 7. Quotation (> )
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            parse_bold_runs(p, stripped[2:], is_italic=True, font_color=RGBColor(102, 102, 102))
            idx += 1
            continue
            
        # 8. Unordered list (- or * )
        if stripped.startswith('- ') or stripped.startswith('* '):
            content_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)
            parse_bold_runs(p, content_text)
            idx += 1
            continue
            
        # 9. Ordinary body paragraph
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(8)
        parse_bold_runs(p, line)
        idx += 1
        
    # If file ends while still in table block
    if in_table:
        render_table(doc, table_lines)
        
    doc.save(docx_path)
    print(f"Successfully created: {docx_path}")

def main():
    md_files = glob.glob(os.path.join(PORTFOLIO_DIR, 'report_*.md'))
    if not md_files:
        print("No report_*.md files found.")
        return
        
    for md_path in md_files:
        filename = os.path.basename(md_path)
        base_name = os.path.splitext(filename)[0] # e.g. report_stt-trainer
        docx_path = os.path.join(PORTFOLIO_DIR, f'{base_name}.docx')
        
        print(f"Processing {filename}...")
        build_single_docx(md_path, docx_path, base_name)

if __name__ == '__main__':
    main()
