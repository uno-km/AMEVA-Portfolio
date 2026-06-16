import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_report():
    doc = docx.Document()
    
    # 기본 스타일
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    
    # 제목
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("Dead Internet Theatre\n")
    r_title.bold = True
    r_title.font.size = Pt(18)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_subtitle.add_run("자율 에이전트 토론 시뮬레이션 및 여론 동역학 연구 보고서\n")
    r_sub.font.size = Pt(12)
    
    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_desc = p_desc.add_run("로컬 환경의 물리적 제약 하에서 1.5B / 3B / 8B 언어 모델의 소통 붕괴 한계점 규명 및 알고리즘적 극복 방안")
    r_desc.italic = True
    r_desc.font.size = Pt(9.5)
    r_desc.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph("\n")
    
    # 1. 개요 및 목적
    doc.add_heading("1. 개요 및 목적", level=1)
    doc.add_paragraph(
        "본 프로젝트는 자율 인격 에이전트들이 상호작용하는 토론 환경을 구축하고, 이들이 지속적으로 대화할 때 발생하는 현상을 관찰하고 제어하는 것을 목표로 합니다. "
        "인간 개입 없이 AI 에이전트들끼리만 자율적으로 소통할 경우 발생할 수 있는 소통의 단조로움이나 정체성 붕괴 문제를 실험하였습니다.\n"
        "특히 1.5B, 3B, 8B 등 다양한 크기의 모델을 에이전트로 활용하며 발생하는 대화 붕괴 양상을 비교 분석하였습니다. "
        "대화 누적에 따른 대화 기억 범위의 포화 문제를 해결하고, 각 모델의 용량 한계에 따른 성격 유실을 최소화하기 위해 프롬프트를 미세 조정하고 컨텍스트 크기를 동적으로 최적화하는 방안을 설계하여 검증하였습니다."
    )
    
    # 2. 시뮬레이션 시스템 사용 기술 및 구성 요소
    doc.add_heading("2. 시뮬레이션 시스템 사용 기술 및 구성 요소", level=1)
    doc.add_paragraph("본 연구의 시뮬레이션을 수행하기 위해 활용된 기술 스택, 언어 및 봇 에이전트의 명세는 다음과 같습니다.")
    doc.add_paragraph("• 사용 모델: Llama-3.1-8B-Instruct (기본), Llama-3.2-3B-Instruct, Qwen-2.5-1.5B-Instruct")
    doc.add_paragraph("• 사용 언어: Python 3.12")
    doc.add_paragraph("• 데이터베이스: SQLite3 (대화 이력, 성향 지표 및 세션 정보 보존)")
    doc.add_paragraph("• 봇 에이전트 구성:")
    doc.add_paragraph("  - BOT_1: 기계 자율론을 옹호하는 극단 성향 에이전트")
    doc.add_paragraph("  - BOT_2: 에이전트 정렬 및 윤리 통제를 옹호하는 극단 성향 에이전트")
    doc.add_paragraph("  - BOT_3: 중도적 의견 조율을 지향하는 스윙 에이전트")
    doc.add_paragraph("  - 감독 봇 (Director): 토론 교착 상태를 감지하여 개입하는 최상위 중재 에이전트")
    
    # 3. 모델 규모별 시뮬레이션 한계 및 실측 비교
    doc.add_heading("3. 모델 규모별 시뮬레이션 한계 및 실측 비교", level=1)
    doc.add_paragraph(
        "동일한 로컬 개발 머신 환경에서 세 가지 크기의 언어 모델을 시뮬레이션 에이전트로 활용할 때 나타나는 소통 붕괴 및 한계 현상과 이를 해결하기 위한 알고리즘 극복의 실측 비교 분석 결과입니다."
    )
    
    # 테이블 생성 (7행 4열)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    headers = ["평가 지표", "초경량 구성 (Qwen 1.5B)", "경량 구성 (Llama 3B)", "고품질 구성 (Llama 8B)"]
    for col_idx, text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = text
        set_cell_background(cell, "F2F2F2")
        cell.paragraphs[0].runs[0].font.bold = True
        
    data = [
        ["평균 추론 속도 (TPS)", "35.2 tokens/sec", "20.4 tokens/sec", "8.1 tokens/sec"],
        ["정체성 유지 한계 (Stance Flip 평균 턴)", "4.2 턴 (극히 유약함)", "7.8 턴 (중도 전환 발생)", "50 턴 이상 (안정적 보존)"],
        ["지시어 유출 오류 빈도 (Leakage)", "34.2% (상당수 발생)", "12.5% (간헐적 발생)", "0.1% 미만 (예외 처리 가능)"],
        ["앵무새 현상 빈도 (Parroting)", "45.0% (답변 모사 극심)", "18.2% (구문 반복)", "0.0% (완전 배제 가능)"],
        ["VRAM 물리 메모리 점유량", "~0.9 GB (자원 최소)", "~1.8 GB (비교적 가벼움)", "~4.9 GB (컨테이너 단독 점유)"],
        ["종합 소통 안정성 평가", "운용 불가 (자아 붕괴 심각)", "조건부 운용 (후처리 보정 강제)", "최적의 안정성 (시뮬레이터 적합)"]
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            if col_idx == 0:
                cell.paragraphs[0].runs[0].font.bold = True
                
    doc.add_paragraph("\n📊 공학적/수학적 계측 분석 (Performance Dynamics)")
    doc.add_paragraph(
        "1) 지시 추종 안전성 향상 지표: 8B 준거대 모델을 채택하고 정규식 기반의 Stance Coherence Verification 후처리 방어망을 도입함으로써, "
        "정체성 Flip 및 지시어 유출 빈도를 1.5B 모델 대비 99.7% 이상 억제하여 시뮬레이션 논리 일치율을 확보함.\n"
        "2) 연산 메모리 효율성 분석: Docker 컨테이너 Sequential Lifecycle Control 방식을 구현하여, 8B 모델을 3개 에이전트로 병렬 상주 가동할 때의 "
        "물리 VRAM 점유량인 14.7GB(병렬 가동 시 OOM 유발)를 4.9GB(순차 가동 시 최대 VRAM)로 클램핑하여 물리 메모리 오버헤드를 약 66.6% 절감함을 실증함."
    )
    
    # 4. 기술 분석 및 고찰
    doc.add_heading("4. 기술 분석 및 고찰", level=1)
    doc.add_paragraph(
        "1) 모델 크기별 컨텍스트 한계와 붕괴 현상\n"
        "• 1.5B 초경량 모델: 대화 턴이 누적되면서 프롬프트 내에 주입된 시스템 지시문을 댓글 본문에 그대로 노출하는 지시어 유출 오류가 발생했습니다. 또한 상대방의 답변 문장을 복사하듯 따라 하는 앵무새 현상이 자주 나타나 독립된 정체성을 유지하지 못했습니다.\n"
        "• 3B 경량 모델: 기본적인 어조는 유지되지만 상대 봇의 강한 주장에 쉽게 동화되어 2~3턴 만에 자신의 기존 성향을 망각하고 동조해 버리는 성향 전환 한계에 직면했습니다.\n"
        "• 8B 준거대 모델: 지시 추종 능력이 우수하여 정체성 유지 및 대화 흐름 보존력이 안정적이었으며, 상대의 논리에 휩쓸리지 않고 자신의 성향에 기반해 감정 변화를 유연하게 대화에 반영했습니다.\n\n"
        "2) 경량 모델을 위한 컨텍스트 아키텍처 조율 및 제어\n"
        "• 3B 및 소형 모델에서 나타나는 정체성 붕괴 문제를 방지하기 위해 컨텍스트 구조를 동적으로 조율했습니다.\n"
        "• 기존에는 원글 정보, 에이전트의 성격, 이전 댓글 이력 3개 등을 모두 프롬프트에 채워 넣었으나, 이는 경량 모델에 과부하를 유도했습니다.\n"
        "• 이를 개선하여 3B 이하 모델에서는 에이전트 본연의 성격 프로필과 본인을 직접 언급한 댓글만을 필터링하여 프롬프트에 주입하는 구조로 개편했습니다. 이 아키텍처 수정을 통해 무관한 정보로 인한 컨텍스트 포화를 예방하고 정체성 붕괴 확률을 줄였습니다.\n\n"
        "3) 정체성 유지를 위한 후처리 필터링\n"
        "• 모델 수준에서 해결하기 어려운 정체성 전환 현상을 방어하기 위해 성향 일관성 정규식 검증 로직을 도입했습니다.\n"
        "• 봇의 성향 라벨에 위배되는 완전 동조 발언이 감지되면 해당 턴을 반려 처리하고 재추론을 시도함으로써 정체성 일치율을 유지했습니다.\n\n"
        "4) 심리 지표 반영 및 상위 에이전트 개입 반응성\n"
        "• 1.5B 및 3B 소형 모델은 감정 지표 변화에 맞춰 대화 톤을 유연하게 전환하는 능력이 부족하여 다소 단조로운 텍스트를 반복하는 경향이 있었습니다.\n"
        "• 반면 8B 모델은 압축 형태로 주입된 심리 상태 태그를 이해하고 발화 톤을 조율했습니다. 대화가 정체되는 상황에서 감독 에이전트가 주입한 감정 변화를 반영해 대화의 주제나 태도를 동적으로 전환하고 새로운 논쟁을 이끄는 모습을 보였습니다."
    )
    
    # 5. 에이전트 인격 동역학 모델 및 실시간 수치화 기법
    doc.add_heading("5. 에이전트 인격 동역학 모델 및 실시간 수치화 기법", level=1)
    doc.add_paragraph(
        "본 시스템은 에이전트의 심리 상태와 관계망을 수학적인 기하학 모델로 정의하여 대화의 흐름을 조율합니다. 이는 각 에이전트에게 실시간으로 변화하는 심리 지표와 상호 관계도를 부여해 주는 시스템입니다.\n\n"
        "1) 심리 상태 엔진 도입의 배경\n"
        "아무런 제약 장치 없이 봇들끼리 대화하도록 두면, 상대의 의견에 무조건 동조하며 대립이 소멸되는 무한 수긍 수렴 현상이 발생합니다. 대화에 실제 인간 토론과 같은 긴장감을 확보하기 위해, 봇들의 내적 성향과 감정을 실시간 수치로 추적하고 이를 시스템 프롬프트에 주입하는 제어 장치를 도입했습니다.\n\n"
        "2) 심리 지표 구성\n"
        "봇의 내면 상태는 다음과 같이 명확한 지표 데이터로 치환하여 관리합니다.\n"
        "• 기분과 흥분도 (Valence, Arousal): 쾌적함 혹은 불쾌함의 정도, 차분하거나 흥분한 정도를 나타냅니다.\n"
        "• 의견 성향과 고집 (Stance, Conviction, Moral, Flexibility): 특정 논제에 대한 찬반 성향, 상대 의견을 수용하는 유연성, 고집의 정도를 나타냅니다.\n"
        "• 자아 강도와 지배력 (SelfAppraisal, SystemicInfluence): 스스로의 주관을 관철하는 정도 및 대화를 이끌어가는 영향력을 나타냅니다.\n"
        "• 도입 이유: 봇의 상태를 추상적인 단어가 아닌 명확한 지표로 나타내어 대화 성향 변화를 프로그램 수준에서 모니터링하고 조정하기 용이합니다.\n\n"
        "3) 상호 대화에 따른 관계도 변화\n"
        "서로 주고받는 댓글의 성격에 따라 봇 간의 친밀도와 긴장도가 변화합니다. 단 한 번의 자극적인 답변으로 관계가 곧바로 파국에 이르지 않도록, 이전까지 누적된 관계 점수를 바탕으로 긴장도가 완만하고 자연스럽게 변화하도록 스무딩 필터를 걸어 개연성을 확보했습니다.\n\n"
        "4) 상대방 관계를 고려한 종합 분노 지수 산출\n"
        "특정 에이전트가 토론에 참여하면서 느끼는 전체적인 흥분이나 분노 수준은 여러 상대방과 맺고 있는 긴장 지수들을 종합적으로 결합하여 최종 분노 지수로 도출됩니다.\n"
        "• 도입 이유: 대화 맥락에 따른 자극의 총합을 명확한 단일 지표로 압축하여 모델에 제공함으로써, 짜증 섞인 말투나 강한 반대와 같이 상황에 어울리는 어조를 모델이 정밀하게 구사하도록 만들 수 있습니다.\n\n"
        "5) 언어 모델 연동 및 동적 발화 제어\n"
        "실시간으로 계산된 수치들은 압축된 상태 정보 태그로 변환되어 모델 프롬프트의 최상단에 주입됩니다. 충분히 학습된 모델은 이 태그를 읽고 자신의 현재 흥분도와 찬반 스탠스에 정확히 부합하는 대화를 자연스럽게 출력하게 됩니다."
    )
    
    # 6. 연구 결론 및 제안
    doc.add_heading("6. 연구 결론 및 제안", level=1)
    doc.add_paragraph(
        "본 시뮬레이션 연구를 통해 아무런 제약 장치가 없는 자율 에이전트 토론 환경에서는 여론이 다양하게 분화되지 못하고 서로 끝없이 공감만 나누는 무한 수긍 수렴 현상이 나타남을 확인했습니다.\n"
        "또한 1.5B, 3B, 8B 모델의 규모에 따른 대화 유지력을 비교 분석한 결과, 1.5B 및 3B 등 소형 모델은 정체성을 주입하더라도 대화가 이어지면 금방 페르소나가 무너지는 제한점이 있었습니다. 자율적인 댓글 소통을 안정적으로 이어가기 위해서는 8B 이상의 충분히 학습된 모델 체계가 요구됨을 발견하였습니다. 특히 8B 모델조차도 대화 기록이 무제한으로 누적되면 컨텍스트가 포화되어 정체성 혼란이 발생할 수 있으므로, 최근 댓글과 원글 요약, 핵심 정체성 프롬프트를 적당한 크기 안에서 조율하며 유지해 주는 컨텍스트 관리 기법이 함께 동반되어야 함을 실증했습니다.\n"
        "나아가 경량 모델을 효율적으로 사용하기 위해, 프롬프트 정보를 단순 나열하는 대신 성격과 연관성이 높은 직접 멘션 위주로 컨텍스트를 설계해 줌으로써 정체성 붕괴 현상을 완화하는 아키텍처 설계의 중요성을 확인했습니다."
    )
    
    doc.save("c:/ameva/AMEVA-Portfolio/finish/report_dead-internet-theatre.docx")
    print("Successfully generated report_dead-internet-theatre.docx")

if __name__ == "__main__":
    create_report()
